#!/usr/bin/env python3
"""
Generate PLRA Central Payment Gateway — New Client / Service Onboarding Guide
Formal SOP + Integration Guide with sample request/response payloads.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


GREEN = RGBColor(0x00, 0x68, 0x37)
NAVY = RGBColor(0x1A, 0x2B, 0x4A)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x44, 0x44, 0x44)
TABLE_HDR = "D9EAD3"


def set_run(run, size=11, bold=False, italic=False, color=BLACK, font="Calibri", mono=False):
    run.font.name = "Consolas" if mono else font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, italic=False, color=BLACK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, mono=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color, mono=mono)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = GREEN if level == 1 else NAVY
        run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text, size=11, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run(r, size=size, bold=True)
        r = p.add_run(text)
        set_run(r, size=size)
    else:
        r = p.add_run(text)
        set_run(r, size=size)
    return p


def numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run(r, size=size)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=9, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=9, color=BLACK)
        shade_cell(cell, TABLE_HDR)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], str(val), size=9)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def page_break(doc):
    doc.add_page_break()


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "006837")
    pBdr.append(bottom)
    pPr.append(pBdr)


def kv_table(doc, pairs):
    add_table(doc, ["Attribute", "Value"], [[k, v] for k, v in pairs], col_widths=[5.5, 11.5])


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.25)
    run = p.add_run(text)
    set_run(run, size=8.5, mono=True, color=NAVY)


def set_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(
            "PLRA-PGW-ONB-GUIDE-001  |  Version 1.1  |  Restricted  |  "
            "Central Payment Gateway — New Client / Service Onboarding"
        )
        set_run(run, size=8, color=GRAY)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Punjab Land Records Authority (PLRA)  |  Govt of the Punjab  |  Page ")
        set_run(run, size=8, color=GRAY)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run2 = fp.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instr)
        run2._r.append(fldChar2)
        set_run(run2, size=8, color=GRAY)


def build():
    doc = Document()
    set_header_footer(doc)

    # ===================== COVER =====================
    add_para(doc, "STANDARD OPERATING PROCEDURE & INTEGRATION GUIDE", size=11, bold=True,
             color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(
        doc,
        "Central Payment Gateway — Onboarding a New Client / Service",
        size=16, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )
    add_para(
        doc,
        "Department / Service Integration Guide for Token, Fetch & Intimate APIs",
        size=12, bold=False, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
    )
    add_hr(doc)

    kv_table(doc, [
        ("Document Title", "PLRA Central Payment Gateway — New Client / Service Onboarding Guide"),
        ("Document Identifier", "PLRA-PGW-ONB-GUIDE-001"),
        ("Version", "1.1"),
        ("Document Status", "Approved for External Issue (Template Instantiation)"),
        ("Security Classification", "Restricted"),
        ("Date of Issue", "20-Aug-2026"),
        ("Issuing Authority", "Punjab Land Records Authority (PLRA) — Payment Gateway Programme"),
        ("Audience", "New client departments / services that own challans or fees and wish to integrate with PLRA Central Payment Gateway"),
        ("Prepared By", "PLRA Payment Gateway / Integration Team"),
        ("Document Owner", "Director IT / Payment Gateway Custodian, PLRA"),
    ])

    add_para(
        doc,
        "This document is the property of the Punjab Land Records Authority (PLRA), Government of the Punjab. "
        "It is issued to enable a named client department or service to design, develop, test and operate "
        "Token, Fetch and Intimate APIs for integration with the PLRA Central Payment Gateway. "
        "It shall not be reproduced or disclosed to any third party without prior written authorisation "
        "of the Document Owner. No live credentials, passwords or production secrets shall be embedded "
        "in this guide — credentials are exchanged only through a secure channel.",
        size=9, italic=True, color=GRAY, space_before=8,
    )

    page_break(doc)

    # ===================== TOC =====================
    heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Introduction and Purpose",
        "2. Payment Gateway Solution Overview",
        "3. Scope and Current Source Systems",
        "4. Product Functions",
        "5. Roles and Responsibilities",
        "6. Standard Operating Procedure (SOP) — Onboarding a New Client",
        "7. Integration Architecture and Runtime Flow",
        "8. Authentication Pattern (Pattern A — Mandatory for New Services)",
        "9. API 1 — Token (Sample Request / Response)",
        "10. API 2 — Fetch / Retrieve Challan (Sample Request / Response)",
        "11. API 3 — Intimate / Payment Confirmation (Sample Request / Response)",
        "12. What PLRA Gateway Does with Your Data",
        "13. Network & Security — Whitelisting (Both Sides)",
        "14. Pre-Go-Live Information Exchange Checklist",
        "15. UAT Checklist (Department / Client)",
        "16. Live / Production Checklist",
        "17. Operational Expectations",
        "18. Mapping from Legacy Fard & eStamp (Transition Only)",
        "19. Frequently Asked Questions",
        "20. Appendix A — Onboarding Message Template to PLRA",
        "21. Appendix B — Document Control",
        "22. Appendix C — Acceptance and Sign-Off",
    ]
    for item in toc_items:
        add_para(doc, item, size=11, space_after=2)

    page_break(doc)

    # ===================== 1 =====================
    heading(doc, "1. Introduction and Purpose", 1)
    add_para(
        doc,
        "This guide is for systems that own challans or fees. The client department exposes three APIs; "
        "the PLRA Central Payment Gateway calls those APIs to:",
    )
    numbered(doc, "Obtain an access token.")
    numbered(doc, "Retrieve unpaid challan / fee details.")
    numbered(doc, "Intimate (confirm) payment after successful card payment.")

    add_para(
        doc,
        "The client does not call MPGS (Mastercard Payment Gateway Services). The client does not call "
        "the Bank of Punjab (BoP) for this flow. PLRA Gateway owns payment collection, fee bifurcation, "
        "BoP intimation, and Intimate retries.",
        space_before=6,
    )
    add_para(
        doc,
        "The purpose of this document is to provide a complete, formal Standard Operating Procedure (SOP) "
        "and technical contract so that a new client can onboard onto the Central Payment Gateway in a "
        "controlled, auditable and repeatable manner — including mandatory JSON payloads, authentication, "
        "network whitelisting, UAT and Live checklists.",
    )

    # ===================== 2 =====================
    heading(doc, "2. Payment Gateway Solution Overview", 1)
    add_para(
        doc,
        "The Payment Gateway solution consists of three integrated components:",
    )
    add_table(
        doc,
        ["#", "Component", "Description"],
        [
            ["1", "Citizen-facing Payment Gateway Web Portal",
             "Allows a citizen or applicant to enter a challan number, fetch challan details, and pay the associated fee online."],
            ["2", "Punjab Zameen Mobile Application Integration",
             "Same payment gateway functionality within Punjab Zameen — in the Record Copy tab and via a dedicated Payment card on the main menu."],
            ["3", "Admin Portal",
             "Allows internal staff to monitor, verify, reconcile and audit payment transactions, inquiries, system logs, service configuration, merchant settings, partners, users/roles and settlement."],
        ],
        col_widths=[1.2, 5.0, 10.5],
    )

    # ===================== 3 =====================
    heading(doc, "3. Scope and Current Source Systems", 1)
    add_para(
        doc,
        "The Payment Gateway Project provides a centralised, secure and auditable online payment mechanism "
        "for government service fees that were previously collected through manual or bank-challan-based processes.",
    )
    add_para(doc, "The system currently supports fee collection for two source systems:", space_before=4)
    bullet(doc, "CLRMIS (Computerized Land Records Management Information System) — Fard Fee (Copy Fee of Fard), Mutation Fee, Crop Inspection Fee, and Registry Fee.",
           bold_prefix="CLRMIS — ")
    bullet(doc, "covering the fee of all E-Stamp instruments / papers.",
           bold_prefix="E-Stamp — ")

    add_para(
        doc,
        "The Gateway identifies the originating (source) system of an incoming challan by reading a prefix "
        "embedded in the challan number, which uniquely identifies whether the challan belongs to CLRMIS, "
        "E-Stamp, or a newly registered service. Once identified, the Gateway calls the relevant source "
        "system’s API to fetch challan details and, after successful payment, calls the relevant API to "
        "intimate / confirm payment back to the source system.",
        space_before=6,
    )
    add_para(
        doc,
        "The Gateway also integrates with the Bank of Punjab (BOP) API for intimation purposes. Because "
        "the service fee is segregated across multiple government account heads, the Gateway sends the "
        "segregated fee breakup to BOP so that the correct proportion of the collected amount is credited "
        "/ intimated against each account head.",
    )

    # ===================== 4 =====================
    heading(doc, "4. Product Functions", 1)
    add_para(doc, "At a high level, the system performs the following functions:")
    funcs = [
        "Accept a challan number as input from a citizen (web or mobile) and identify the source system using the challan number prefix.",
        "Call the appropriate source system API to fetch challan details, including the payer’s Name, CNIC, and Challan Fee amount.",
        "Display challan details to the citizen and indicate whether the challan is Paid or Unpaid.",
        "Enable a “Pay Now” action for unpaid challans, and display a “Paid” indicator for already-paid challans.",
        "Process payment through a debit / credit card payment flow (MPGS).",
        "Segregate the collected fee into the applicable government account heads and intimate the segregated breakup to BOP via the BOP API.",
        "Confirm / intimate successful payment back to the originating source system via API.",
        "Record every challan-fetch attempt as an Inquiry, and every verified successful payment as a Payment Ledger entry.",
        "Provide an Admin Portal for dashboard monitoring, inquiry tracking, ledger review, service/API configuration, gateway/merchant settings, integration partner management, user and role administration, audit and integration logging, and settlement reconciliation.",
    ]
    for f in funcs:
        bullet(doc, f)

    page_break(doc)

    # ===================== 5 =====================
    heading(doc, "5. Roles and Responsibilities", 1)
    add_table(
        doc,
        ["Party", "Responsibility"],
        [
            ["Client Department / Service",
             "Expose Token, Fetch and Intimate APIs (Pattern A); supply UAT/Prod URLs and credentials; implement mandatory JSON; support UAT/Live; keep Intimate idempotent; whitelist PLRA egress IPs if APIs are IP-restricted."],
            ["PLRA Payment Gateway",
             "Citizen payment UX (web + Punjab Zameen); MPGS card collection; fee bifurcation; BOP intimation; call client Token/Fetch/Intimate; registry/prefix assignment; outbound allowlisting; retries on Intimate failure; Admin Portal operations."],
            ["Bank of Punjab (BOP)",
             "Receive segregated fee breakup intimation from PLRA Gateway (client does not call BOP for this flow)."],
            ["MPGS",
             "Card payment processing owned by PLRA Gateway (client does not call MPGS)."],
        ],
        col_widths=[4.5, 12.2],
    )

    add_para(doc, "What the client SHALL NOT do:", size=11, bold=True, space_before=4)
    bullet(doc, "Call MPGS for this payment flow.")
    bullet(doc, "Call BoP for bifurcation / settlement of Gateway-collected fees.")
    bullet(doc, "Invent or self-assign a challan prefix — PLRA assigns the prefix at registration.")
    bullet(doc, "Expose Token credentials in mobile apps, public repositories or uncontrolled channels.")
    bullet(doc, "Use static long-lived headers (legacy Pattern B) for new integrations unless PLRA approves a written exception.")

    # ===================== 6 SOP =====================
    heading(doc, "6. Standard Operating Procedure (SOP) — Onboarding a New Client", 1)
    add_para(
        doc,
        "The following SOP is mandatory for onboarding any new client department or service onto the "
        "Central Payment Gateway. Steps shall be completed in sequence unless PLRA formally agrees otherwise.",
    )

    heading(doc, "6.1 SOP Summary", 2)
    add_table(
        doc,
        ["Step", "Activity", "Owner", "Output"],
        [
            ["1", "Raise onboarding request to PLRA Payment Gateway programme", "Client", "Onboarding ticket / formal request"],
            ["2", "PLRA creates RegistryServices entry and assigns 3-digit challan prefix", "PLRA", "Prefix (e.g. 003) + registry record"],
            ["3", "Client delivers UAT Token / Fetch / Intimate URLs + credentials (secure channel)", "Client", "UAT endpoints + secrets"],
            ["4", "Exchange sample unpaid / already-paid / invalid references", "Client", "Sample reference set"],
            ["5", "Network allowlisting — both sides (see Section 13)", "Both", "Outbound + inbound firewall rules"],
            ["6", "Joint connectivity test: Gateway → Token → Fetch", "Both", "Connectivity evidence"],
            ["7", "Joint UAT per Section 15 (unpaid, paid, not-found, intimate, idempotent retry)", "Both", "Signed UAT checklist"],
            ["8", "Defect closure / formal deferral", "Both", "UAT sign-off"],
            ["9", "Production cutover plan: separate Prod URLs, secrets, IP allowlists", "Both", "Prod readiness record"],
            ["10", "Live checklist (Section 16) and controlled go-live + hypercare", "Both", "Go-live report"],
        ],
        col_widths=[1.2, 7.5, 2.5, 5.5],
    )

    heading(doc, "6.2 Detailed SOP Steps", 2)

    heading(doc, "Step 1 — Request Registration", 3)
    bullet(doc, "Client sends the onboarding message in Appendix A to PLRA.")
    bullet(doc, "PLRA acknowledges and opens a controlled onboarding ticket.")

    heading(doc, "Step 2 — Registry & Prefix Assignment", 3)
    bullet(doc, "PLRA creates a RegistryServices entry for the client service.")
    bullet(doc, "PLRA assigns a unique 3-digit challan prefix (examples: 001 = Fard, 002 = eStamp, 003 = next new service).")
    bullet(doc, "Client SHALL NOT invent or reuse another service’s prefix.")

    heading(doc, "Step 3 — Deliver UAT Integration Assets", 3)
    bullet(doc, "Token URL + grant type + username/password (or agreed OAuth fields).")
    bullet(doc, "Fetch URL template (with {reference} or ChallanNumber= placeholder).")
    bullet(doc, "Intimate URL.")
    bullet(doc, "Confirmation that Fetch and Intimate use Authorization: Bearer {access_token} (Pattern A).")
    bullet(doc, "Whether client APIs are IP-restricted (Yes/No).")

    heading(doc, "Step 4 — Sample References", 3)
    bullet(doc, "Provide at least one unpaid reference, one already-paid reference, and one invalid / not-found reference.")

    heading(doc, "Step 5 — Dual-Side Whitelisting", 3)
    bullet(doc, "PLRA allowlists client UAT host on Gateway outbound allowlist (Side A).")
    bullet(doc, "If client APIs are IP-locked, client allowlists PLRA UAT egress IP(s) (Side B).")
    bullet(doc, "No live call is expected to succeed until both sides confirm allowlisting.")

    heading(doc, "Step 6–7 — Joint Connectivity & UAT", 3)
    bullet(doc, "Execute Section 15 UAT checklist end-to-end.")
    bullet(doc, "Mandatory scenarios: unpaid fetch → pay → intimate success; already-paid fetch; not-found; duplicate intimate idempotency; invalid/expired Bearer.")

    heading(doc, "Step 8–10 — Sign-off, Production & Live", 3)
    bullet(doc, "UAT sign-off precedes any production credential or prod IP issuance.")
    bullet(doc, "Production uses separate URLs, separate secrets, and separate allowlists (Section 16).")
    bullet(doc, "Controlled go-live with hypercare and named contacts for failed intimations.")

    page_break(doc)

    # ===================== 7 =====================
    heading(doc, "7. Integration Architecture and Runtime Flow", 1)
    add_para(doc, "Runtime flow (PLRA Gateway → Client Service):", size=11, bold=True)
    code_block(
        doc,
        "Citizen pays via PLRA Gateway (Web Portal / Punjab Zameen)\n"
        "        |\n"
        "        v\n"
        "1) Gateway -> POST  Client Token API     -> access_token\n"
        "2) Gateway -> GET   Client Fetch API     + Bearer -> standard fetch JSON\n"
        "        |\n"
        "   (citizen pays on MPGS)\n"
        "        |\n"
        "        v\n"
        "3) Gateway -> POST  Client Intimate API  + Bearer + standard intimate JSON\n"
        "\n"
        "Parallel (owned by PLRA Gateway, not by client):\n"
        "   - MPGS card capture\n"
        "   - Fee bifurcation using feeDetails[].accountNo / accountTitle\n"
        "   - BoP intimation of segregated account-head breakup\n"
        "   - Intimate retries on failure",
    )

    add_para(doc, "Mandatory APIs the client must provide:", size=11, bold=True, space_before=4)
    add_table(
        doc,
        ["#", "API", "Method", "Purpose"],
        [
            ["1", "Token", "POST", "Issue short-lived access token for Gateway"],
            ["2", "Fetch (Retrieve Challan)", "GET", "Return challan / fee / applicant / territory by reference"],
            ["3", "Intimate (Payment Confirmation)", "POST", "Mark challan paid after successful MPGS payment"],
        ],
        col_widths=[1.2, 5.5, 2.5, 7.5],
    )
    add_para(
        doc,
        "Auth for Fetch and Intimate: Authorization: Bearer {access_token from Token API}. "
        "This is Pattern A (same model as eStamp). New services MUST follow Pattern A. "
        "Static long-lived headers (current Fard style) are legacy Pattern B — not accepted for new "
        "integrations unless PLRA approves a written exception.",
        size=10,
    )

    # ===================== 8 =====================
    heading(doc, "8. Authentication Pattern (Pattern A — Mandatory for New Services)", 1)
    add_table(
        doc,
        ["Item", "Requirement"],
        [
            ["Pattern", "Pattern A — OAuth-style Token API + Bearer on Fetch/Intimate"],
            ["Token caching", "Gateway caches access_token and reuses until near expiry, then refreshes"],
            ["Per-request token", "Prohibited — client Token API would be overloaded"],
            ["Secrets custody", "Username/password (or private key material) stored in vault only"],
            ["Fetch / Intimate auth header", "Authorization: Bearer {access_token}"],
            ["Legacy Pattern B", "Static long-lived headers — exception only with PLRA approval"],
        ],
        col_widths=[5.0, 11.7],
    )

    page_break(doc)

    # ===================== 9 TOKEN =====================
    heading(doc, "9. API 1 — Token", 1)
    heading(doc, "9.1 Purpose", 2)
    add_para(
        doc,
        "Gateway calls the client Token API, caches the access token, and uses it on Fetch / Intimate "
        "until near expiry (then refreshes).",
    )

    heading(doc, "9.2 Recommended Contract (eStamp Reference — Pattern A)", 2)
    add_para(doc, "Request", size=11, bold=True)
    code_block(
        doc,
        "POST {TokenUrl}\n"
        "Content-Type: application/x-www-form-urlencoded\n"
        "Accept: application/json\n"
        "\n"
        "Body (form fields):\n"
        "  grant_type=password\n"
        "  userName={issued-to-PLRA}\n"
        "  password={issued-to-PLRA — vault only}",
    )

    add_table(
        doc,
        ["Field", "Example / Rule"],
        [
            ["grant_type", "password (or other grant agreed at onboarding)"],
            ["userName", "Issued to PLRA for this service / environment"],
            ["password", "Issued to PLRA — vault only; never in this document or public repos"],
        ],
        col_widths=[4.0, 12.7],
    )

    add_para(doc, "Success response (sample)", size=11, bold=True, space_before=4)
    code_block(
        doc,
        "HTTP/1.1 200 OK\n"
        "Content-Type: application/json\n"
        "\n"
        "{\n"
        '  "access_token": "eyJhbGciOi...",\n'
        '  "token_type": "bearer",\n'
        '  "expires_in": 3600\n'
        "}",
    )

    add_table(
        doc,
        ["Field", "Required", "Notes"],
        [
            ["access_token", "Yes", "Bearer token used on Fetch and Intimate"],
            ["token_type", "Recommended", "Typically bearer"],
            ["expires_in", "Recommended", "Seconds. If omitted, Gateway assumes a safe default (~3600)"],
        ],
        col_widths=[4.0, 3.0, 9.7],
    )

    add_para(doc, "Error behaviour:", size=11, bold=True)
    bullet(doc, "Non-2xx → Gateway cannot fetch/intimate; treat as client outage.")
    bullet(doc, "Invalid credentials → Gateway raises operational alert; do not fall back to unauthenticated calls.")

    # ===================== 10 FETCH =====================
    heading(doc, "10. API 2 — Fetch (Retrieve Challan)", 1)
    heading(doc, "10.1 Purpose", 2)
    add_para(
        doc,
        "Return everything Gateway needs to show the bill, create an Inquiry, charge the citizen, "
        "and later bifurcate fees to BoP.",
    )

    heading(doc, "10.2 Request (Gateway → Client)", 2)
    code_block(
        doc,
        "GET {FetchUrl}?ChallanNumber={reference}\n"
        "Authorization: Bearer {access_token}\n"
        "Accept: application/json",
    )
    bullet(doc, "{reference} is the challan without the PLRA prefix (Gateway strips 001 / 002 / … before calling you).")
    bullet(doc, "URL template may use {reference} or {challan_number} — agree the placeholder with PLRA at onboarding.")
    bullet(doc, "Method is typically GET.")

    heading(doc, "10.3 Response — Mandatory Standard JSON (Client → Gateway)", 2)
    add_para(doc, "Sample — Unpaid, payable challan (HTTP 200):", size=11, bold=True)
    code_block(
        doc,
        "HTTP/1.1 200 OK\n"
        "Content-Type: application/json; charset=utf-8\n"
        "\n"
        "{\n"
        '  "transactionInfo": {\n'
        '    "applicantName": "ALI KHAN",\n'
        '    "cnic": "3520212345671",\n'
        '    "transactionNo": "89439714",\n'
        '    "description": "Mutation Fee",\n'
        '    "totalFee": 3500.00,\n'
        '    "isPaid": false,\n'
        '    "consumerNumber": "",\n'
        '    "deptTransactionId": "",\n'
        '    "territory": {\n'
        '      "divisionId": null,\n'
        '      "divisionName": "",\n'
        '      "districtId": 42,\n'
        '      "districtName": "LAHORE",\n'
        '      "tehsilId": null,\n'
        '      "tehsilName": "LAHORE CITY",\n'
        '      "mauzaId": null,\n'
        '      "mauzaName": "SOME MAUZA"\n'
        "    }\n"
        "  },\n"
        '  "feeDetails": [\n'
        "    {\n"
        '      "feeTypeId": 10,\n'
        '      "feeTypeName": "Service Fee",\n'
        '      "accountNo": "B01418",\n'
        '      "accountTitle": "PLRA COLLECTION ACCOUNT",\n'
        '      "fee": 100.00\n'
        "    },\n"
        "    {\n"
        '      "feeTypeId": 11,\n'
        '      "feeTypeName": "Government Fee",\n'
        '      "accountNo": "6580036040700020",\n'
        '      "accountTitle": "PLRA COLLECTION ACCOUNT",\n'
        '      "fee": 3400.00\n'
        "    }\n"
        "  ]\n"
        "}",
    )

    add_para(doc, "Sample — Already paid (HTTP 200):", size=11, bold=True, space_before=4)
    code_block(
        doc,
        "{\n"
        '  "transactionInfo": {\n'
        '    "applicantName": "ALI KHAN",\n'
        '    "cnic": "3520212345671",\n'
        '    "transactionNo": "89439714",\n'
        '    "description": "Mutation Fee",\n'
        '    "totalFee": 3500.00,\n'
        '    "isPaid": true,\n'
        '    "consumerNumber": "",\n'
        '    "deptTransactionId": "",\n'
        '    "territory": {\n'
        '      "divisionId": null,\n'
        '      "divisionName": "",\n'
        '      "districtId": 42,\n'
        '      "districtName": "LAHORE",\n'
        '      "tehsilId": null,\n'
        '      "tehsilName": "LAHORE CITY",\n'
        '      "mauzaId": null,\n'
        '      "mauzaName": "SOME MAUZA"\n'
        "    }\n"
        "  },\n"
        '  "feeDetails": [ /* same structure; accounts still present */ ]\n'
        "}",
    )

    add_para(doc, "Sample — Not found / invalid (preferred HTTP 404):", size=11, bold=True, space_before=4)
    code_block(
        doc,
        "HTTP/1.1 404 Not Found\n"
        "Content-Type: application/json; charset=utf-8\n"
        "\n"
        "{\n"
        '  "status": false,\n'
        '  "message": "Challan not found for the supplied reference"\n'
        "}",
    )

    page_break(doc)

    heading(doc, "10.4 Field Rules — transactionInfo", 2)
    add_table(
        doc,
        ["Field", "Required?", "Notes"],
        [
            ["applicantName", "Yes", "Full name"],
            ["cnic", "Yes", "13 digits preferred (dashes optional). Needed for BoP"],
            ["transactionNo", "Yes", "Your challan / PSID / serial (stored as challan ref)"],
            ["description", "Yes", "Short label on receipt (e.g. “E-stamp”, “Fard - Mutation”). Do not put full address here"],
            ["totalFee", "Yes", "Sum of your fee lines (exclude PLRA IPG/MDR)"],
            ["isPaid", "Yes", "true if already paid in your system"],
            ["consumerNumber", "Optional", "Empty \"\" if N/A"],
            ["deptTransactionId", "Optional", "Empty \"\" if N/A"],
            ["territory", "Yes", "Object — see Section 10.5"],
        ],
        col_widths=[4.5, 2.5, 9.7],
    )

    heading(doc, "10.5 Field Rules — territory (Mandatory Object)", 2)
    add_para(
        doc,
        "Send id + name for every level you know. Unknown → null id and \"\" name. "
        "Gateway / BoP must resolve division + district.",
    )
    add_table(
        doc,
        ["Minimum Combination", "Required Fields"],
        [
            ["District-based (eStamp-like)", "districtId and districtName"],
            ["Mauza/tehsil-based (Fard-like)", "mauzaName and tehsilName (ids if you have them)"],
            ["Ideal (all new services)", "District + tehsil + mauza with id + name where known"],
        ],
        col_widths=[6.0, 10.7],
    )

    heading(doc, "10.6 Field Rules — feeDetails[]", 2)
    add_table(
        doc,
        ["Field", "Required?", "Notes"],
        [
            ["fee", "Yes", "Line amount"],
            ["accountNo", "Yes", "Collection account — without this, online pay is blocked"],
            ["accountTitle", "Yes", "Account / duty title — without this, online pay is blocked"],
            ["feeTypeName", "Recommended", "Shown / sent to BoP"],
            ["feeTypeId", "Optional", "null if you have no numeric type"],
        ],
        col_widths=[4.5, 2.5, 9.7],
    )

    heading(doc, "10.7 Fetch Outcomes You Must Support", 2)
    add_table(
        doc,
        ["Case", "HTTP", "Body"],
        [
            ["Unpaid, payable", "200", "Standard JSON, isPaid: false, fee lines with accounts"],
            ["Already paid", "200", "Standard JSON, isPaid: true"],
            ["Not found / invalid", "404 (preferred) or agreed not-found signal", "Clear message"],
        ],
        col_widths=[4.5, 5.5, 6.7],
    )
    add_para(
        doc,
        "Do not invent a different schema per service. Optional keys may be empty, but required keys must be present.",
        size=10, italic=True, color=GRAY,
    )

    page_break(doc)

    # ===================== 11 INTIMATE =====================
    heading(doc, "11. API 3 — Intimate (Payment Confirmation)", 1)
    heading(doc, "11.1 Purpose", 2)
    add_para(
        doc,
        "After MPGS success, Gateway tells your system the challan is PAID so you unlock the service "
        "(print fard, issue stamp, release record copy, etc.).",
    )

    heading(doc, "11.2 Request (Gateway → Client)", 2)
    code_block(
        doc,
        "POST {IntimateUrl}\n"
        "Authorization: Bearer {access_token}\n"
        "Content-Type: application/json\n"
        "\n"
        "{\n"
        '  "challanNumber": "89439714",\n'
        '  "consumerNumber": "89439714",\n'
        '  "deptTransactionId": "89439714",\n'
        '  "psidStatus": "PAID",\n'
        '  "amountPaid": "3500",\n'
        '  "paidDate": "2026-08-10",\n'
        '  "paidTime": "15:30:00",\n'
        '  "bankCode": "BOP",\n'
        '  "requestedBy": "PLRA-PaymentGateway",\n'
        '  "applicantName": "ALI KHAN",\n'
        '  "cnic": "3520212345671",\n'
        '  "cashierName": "ALI KHAN",\n'
        '  "branchName": "online",\n'
        '  "branchCode": "",\n'
        '  "organizationUserId": "",\n'
        '  "consumerType": ""\n'
        "}",
    )

    add_table(
        doc,
        ["Field", "Required for you to accept?", "Notes"],
        [
            ["challanNumber", "Yes", "Same as fetch transactionNo"],
            ["amountPaid", "Yes", "Your base fee only — not IPG/MDR"],
            ["psidStatus", "Yes", '"PAID"'],
            ["paidDate / paidTime", "Yes", "Pakistan local wall clock"],
            ["bankCode / requestedBy", "Recommended", "Values agreed at onboarding"],
            ["consumerNumber / deptTransactionId", "If your system needs them", "May equal challan"],
            ["applicantName / cnic / cashierName", "Optional", "May be empty strings"],
            ["branchName / branchCode / organizationUserId / consumerType", "Optional", "Empty if unused"],
        ],
        col_widths=[5.5, 4.0, 7.2],
    )
    add_para(
        doc,
        "Your Intimate API SHOULD ignore unknown extra fields safely (forward compatible).",
        size=10, italic=True, color=GRAY,
    )

    heading(doc, "11.3 Response (Client → Gateway) — Recommended Standard", 2)
    add_para(doc, "Preferred success:", size=11, bold=True)
    code_block(
        doc,
        "{\n"
        '  "status": true,\n'
        '  "message": "OK"\n'
        "}",
    )
    add_para(doc, "Also accepted (legacy numeric style):", size=11, bold=True)
    code_block(
        doc,
        "{\n"
        '  "Status": 1,\n'
        '  "Message": "Success"\n'
        "}",
    )

    add_table(
        doc,
        ["Your Signal", "Gateway Outcome"],
        [
            ["status: true or Status: 1", "SUCCESS"],
            ["Already paid / Status: 2", "ALREADY_PAID (treated as success for retries)"],
            ["PITB paid / Status: 8", "PITB_PAID (treated as success for retries)"],
            ["status: false / other / HTTP error", "FAILED — Gateway retries"],
        ],
        col_widths=[7.0, 9.7],
    )

    add_para(doc, "Idempotency (mandatory):", size=11, bold=True, space_before=4)
    add_para(
        doc,
        "If Gateway calls Intimate twice for the same paid challan, return success / already-paid — "
        "do not fail the second call. Gateway retries are normal after timeouts or transient faults.",
    )

    page_break(doc)

    # ===================== 12 =====================
    heading(doc, "12. What PLRA Gateway Does with Your Data", 1)
    add_para(doc, "So you know why fields matter:", size=10, italic=True, color=GRAY)
    add_table(
        doc,
        ["Your Fetch Field", "Why PLRA Needs It"],
        [
            ["applicantName / cnic", "Receipt + BoP bifurcation"],
            ["territory.*", "BoP division/district resolution"],
            ["feeDetails[].accountNo / accountTitle", "BoP fee split — missing → online pay blocked"],
            ["totalFee / fee lines", "Citizen payable base (IPG added by Gateway)"],
            ["Intimate amountPaid", "Must match department base — not card total with MDR"],
        ],
        col_widths=[6.5, 10.2],
    )
    add_para(
        doc,
        "End-user apps and other systems never call your three APIs for this flow. Only PLRA Gateway does.",
        size=10, bold=True,
    )

    # ===================== 13 NETWORK =====================
    heading(doc, "13. Network & Security — Whitelisting (Both Sides)", 1)
    add_para(
        doc,
        "Direction of traffic: PLRA Gateway calls your APIs. Allowlisting is therefore two-way. "
        "Do not assume only one side needs firewall changes.",
    )
    code_block(
        doc,
        "PLRA Gateway (egress IP)  ------->  Your Token / Fetch / Intimate APIs\n"
        "        |                                    |\n"
        " (1) PLRA must allow outbound                (2) You may need to allow\n"
        "     to your host/URL                            inbound from PLRA IPs\n"
        "        v                                    v\n"
        " Gateway AllowedHosts /                   Your firewall / WAF /\n"
        " DC outbound policy                       API gateway IP allowlist",
    )

    heading(doc, "13.1 Side A — PLRA Whitelists Your UAT (and later Prod) APIs", 2)
    add_para(
        doc,
        "Before Gateway can call you, PLRA registers your hostnames on the Gateway outbound allowlist "
        "(SSRF / AllowedHosts / DC firewall as applicable).",
    )
    add_table(
        doc,
        ["You Provide", "PLRA Does"],
        [
            ["Full UAT Token / Fetch / Intimate URLs (or base host + paths)", "Adds host to Gateway outbound allowlist"],
            ["Later: Production URLs (separate from UAT)", "Separate prod allowlist entry"],
        ],
        col_widths=[8.5, 8.2],
    )
    add_para(
        doc,
        "Until this is done, Gateway calls to your APIs will be blocked by design (even if credentials are correct).",
        size=10, italic=True, color=GRAY,
    )
    add_para(doc, "Typical rule shape (PLRA side):", size=11, bold=True)
    code_block(
        doc,
        "Direction:        Outbound from PLRA Gateway\n"
        "Source:           PLRA Gateway server / VIP egress\n"
        "Destination:      <YOUR_UAT_API_HOST>\n"
        "Destination Port: 443 (HTTPS) or agreed port\n"
        "Purpose:          Token + Fetch + Intimate",
    )

    heading(doc, "13.2 Side B — You Whitelist PLRA Gateway Egress IPs", 2)
    add_para(
        doc,
        "Many services keep Token/Fetch/Intimate off the open internet and only allow known callers.",
    )
    add_table(
        doc,
        ["If Your APIs Are…", "What To Do"],
        [
            ["Open to internet (auth by token only)", "Still recommended to restrict by IP when possible"],
            ["Already IP-restricted / behind WAF", "You MUST whitelist PLRA’s UAT egress IP(s) or Gateway calls will time out / get 403"],
        ],
        col_widths=[7.0, 9.7],
    )
    add_para(
        doc,
        "PLRA will share UAT egress IP(s) during onboarding and Production egress IP(s) only after UAT sign-off. "
        "You should not guess IPs — ask PLRA in the onboarding ticket.",
    )
    add_para(doc, "Typical rule shape (your side):", size=11, bold=True)
    code_block(
        doc,
        "Direction:        Inbound to your API host\n"
        "Source IP:        <PLRA_GATEWAY_UAT_EGRESS_IP>   <- PLRA provides\n"
        "Source Port:      Any\n"
        "Destination:      <your Token/Fetch/Intimate host>\n"
        "Destination Port: 443 (or agreed)\n"
        "Purpose:          PLRA Payment Gateway -> your 3 APIs",
    )

    heading(doc, "13.3 UAT vs Production", 2)
    add_table(
        doc,
        ["Item", "UAT", "Production"],
        [
            ["Your API URLs", "UAT hosts", "Different prod hosts (preferred)"],
            ["PLRA allowlist of your host", "UAT entry", "New prod entry"],
            ["Your allowlist of PLRA IPs", "UAT egress IPs", "New prod egress IPs"],
            ["Token credentials", "UAT user/password", "Separate prod secrets"],
        ],
        col_widths=[5.0, 5.8, 5.9],
    )
    add_para(
        doc,
        "Do not reuse UAT allowlist rules for production without an explicit prod change request.",
        size=10, bold=True,
    )

    heading(doc, "13.4 Other Security Rules", 2)
    add_table(
        doc,
        ["Topic", "Rule"],
        [
            ["Who calls whom", "Only PLRA Gateway → your 3 APIs (no other channel should call Token/Fetch/Intimate for this payment flow)"],
            ["TLS", "HTTPS in UAT/prod as agreed; valid certificates"],
            ["Secrets", "Token password in vault only — never in mobile apps or public repos"],
            ["URL / IP change", "Notify PLRA before DNS/IP cutover so both allowlists can be updated"],
            ["Timeouts", "Gateway may time out slow intimates and retry — your Intimate must be idempotent"],
        ],
        col_widths=[4.0, 12.7],
    )

    page_break(doc)

    # ===================== 14 =====================
    heading(doc, "14. Pre-Go-Live Information Exchange Checklist", 1)
    add_para(doc, "What PLRA needs from you before go-live:", size=11, bold=True)
    add_table(
        doc,
        ["Item", "Required"],
        [
            ["Token URL + grant type + username/password (or agreed OAuth fields)", "Yes"],
            ["Fetch URL (with {reference} or ChallanNumber= placeholder)", "Yes"],
            ["Intimate URL", "Yes"],
            ["Sample unpaid + already-paid + invalid references", "Yes"],
            ["Confirmation that Fetch/Intimate use Bearer from Token", "Yes"],
            ["Your API hostnames (UAT, later Prod) for PLRA outbound allowlist", "Yes"],
            ["Confirmation whether your firewall IP-restricts callers (so PLRA must send egress IPs)", "Yes"],
            ["Contact for support / failed intimations", "Yes"],
        ],
        col_widths=[13.5, 3.2],
    )

    # ===================== 15 UAT =====================
    heading(doc, "15. UAT Checklist (Department / Client)", 1)
    add_para(
        doc,
        "All items below SHALL be completed and evidenced before UAT sign-off. "
        "Mark each item Pass / Fail with evidence reference.",
    )
    add_table(
        doc,
        ["#", "UAT Check", "Result", "Evidence"],
        [
            ["1", "Token returns access_token + usable expires_in", "☐", ""],
            ["2", "Fetch unpaid → standard JSON, isPaid: false, accounts present", "☐", ""],
            ["3", "Fetch already-paid → isPaid: true", "☐", ""],
            ["4", "Fetch invalid → 404 / agreed not-found", "☐", ""],
            ["5", "Intimate with standard body → success JSON", "☐", ""],
            ["6", "Intimate duplicate → success / already-paid (idempotent)", "☐", ""],
            ["7", "Bearer rejected when token invalid/expired", "☐", ""],
            ["8", "Prefix assigned by PLRA; sample challans shared", "☐", ""],
            ["9", "PLRA has allowlisted your UAT host (Side A — §13.1)", "☐", ""],
            ["10", "You have allowlisted PLRA UAT egress IP(s) if APIs are IP-locked (Side B — §13.2)", "☐", ""],
            ["11", "Joint connectivity test: Gateway → Token → Fetch from PLRA network", "☐", ""],
            ["12", "Prod plan noted: separate URLs + separate IP allowlists after UAT sign-off", "☐", ""],
            ["13", "End-to-end unpaid fetch → citizen pay (MPGS test) → intimate success", "☐", ""],
            ["14", "feeDetails accountNo / accountTitle present (online pay not blocked)", "☐", ""],
            ["15", "Support contact confirmed for failed intimations", "☐", ""],
        ],
        col_widths=[1.2, 11.5, 1.8, 2.2],
    )

    page_break(doc)

    # ===================== 16 LIVE =====================
    heading(doc, "16. Live / Production Checklist", 1)
    add_para(
        doc,
        "Production cutover SHALL NOT proceed until UAT is signed off. "
        "Production assets are separate from UAT.",
    )
    add_table(
        doc,
        ["#", "Live / Production Check", "Result", "Evidence"],
        [
            ["1", "UAT checklist fully signed off with no open critical defects", "☐", ""],
            ["2", "Production Token / Fetch / Intimate URLs delivered (distinct from UAT)", "☐", ""],
            ["3", "Production Token credentials issued via secure channel and stored in vault", "☐", ""],
            ["4", "PLRA has allowlisted Production host on Gateway outbound allowlist", "☐", ""],
            ["5", "Client has allowlisted PLRA Production egress IP(s) (if IP-restricted)", "☐", ""],
            ["6", "Production prefix / registry entry confirmed active", "☐", ""],
            ["7", "Smoke test: Token → Fetch unpaid (prod-safe sample) → Intimate dry-run/agreed path", "☐", ""],
            ["8", "Monitoring / alerting contacts confirmed on both sides", "☐", ""],
            ["9", "Rollback / disable plan agreed (Gateway registry switch / feature flag)", "☐", ""],
            ["10", "Hypercare window and escalation matrix agreed", "☐", ""],
            ["11", "No UAT secrets reused in Production", "☐", ""],
            ["12", "Go-live authorisation recorded by competent authorities of both parties", "☐", ""],
        ],
        col_widths=[1.2, 11.5, 1.8, 2.2],
    )

    # ===================== 17 =====================
    heading(doc, "17. Operational Expectations", 1)
    add_table(
        doc,
        ["Topic", "Expectation"],
        [
            ["Availability", "Token/Fetch/Intimate must be reachable during payment hours agreed with PLRA"],
            ["Intimate failure", "Payment may already be successful at bank; Gateway retries intimate — fix quickly; do not ask citizen to pay again"],
            ["Clock", "paidDate / paidTime use Pakistan local time; keep server clocks reasonably synced"],
            ["Charset", "UTF-8 JSON"],
            ["Content-Type", "Intimate: application/json; Token: form-urlencoded as agreed"],
            ["Breaking changes", "New required fields / URL path changes need coordinated release with PLRA"],
            ["Logging", "Log correlation-friendly fields (challanNumber, HTTP status) for joint troubleshooting — never log full token passwords"],
        ],
        col_widths=[4.0, 12.7],
    )

    # ===================== 18 =====================
    heading(doc, "18. Mapping from Legacy Fard & eStamp (Transition Only)", 1)
    add_para(
        doc,
        "Existing services still use different wire shapes. New services must implement the standard "
        "in Sections 9–11 directly. Fard/eStamp will be migrated onto this same standard.",
    )
    add_table(
        doc,
        ["Standard Field", "Fard Today (Legacy)", "eStamp Today (Closer to Pattern A)"],
        [
            ["Token API", "None (static headers)", "OAuth password grant"],
            ["Fetch auth", "Static headers", "Bearer"],
            ["Intimate", "Query-string POST", "JSON + Bearer"],
            ["applicantName", "applicantName", "TransactionInfo.UserName"],
            ["cnic", "applicantId", "TransactionInfo.CNIC"],
            ["transactionNo", "challanNo", "DeptTransactionId / TransactionNo"],
            ["isPaid", "status == false ⇒ paid", "IsPaid"],
            ["Fee account", "accountNo / accountTitle", "DutyAccountHead / DutyType"],
        ],
        col_widths=[4.0, 6.3, 6.4],
    )

    page_break(doc)

    # ===================== 19 FAQ =====================
    heading(doc, "19. Frequently Asked Questions", 1)
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ["Do we call PLRA payment APIs?", "Not for owning challans. You expose 3 APIs; Gateway calls you."],
            ["Do we need to whitelist anything?", "Yes, if your APIs are IP-restricted — whitelist PLRA Gateway UAT (then prod) egress IPs. PLRA also whitelists your hosts outbound."],
            ["Who starts whitelist?", "Exchange: you send UAT URLs/hosts; PLRA sends UAT egress IPs; both sides apply rules before testing."],
            ["When do we get prod IPs?", "After UAT is finalized / signed off — PLRA issues prod egress IPs and you open prod allowlist separately."],
            ["Can we keep Fard-style static headers?", "Only as approved Pattern B exception. New services = Pattern A."],
            ["Who assigns prefix 003 etc.?", "PLRA, at registration."],
            ["Can fee lines omit account numbers?", "No — online payment will be blocked."],
            ["Does intimate amount include bank MDR?", "No — base department fee only."],
            ["Must every territory id be filled?", "Fill what you have; meet minimum combination for district or mauza+tehsil."],
            ["Intimate called twice?", "Return success / already-paid — Gateway retries are normal."],
        ],
        col_widths=[5.5, 11.2],
    )

    # ===================== Appendix A =====================
    heading(doc, "20. Appendix A — Onboarding Message Template to PLRA", 1)
    add_para(doc, "Client sends:", size=11, bold=True)
    numbered(doc, "Organisation / service name")
    numbered(doc, "Proposed service code (e.g. myservice)")
    numbered(doc, "Token / Fetch / Intimate UAT URLs (hostnames clearly stated)")
    numbered(doc, "Token credentials (secure channel)")
    numbered(doc, "Sample unpaid + paid + invalid references")
    numbered(doc, "Confirmation: Pattern A Bearer on Fetch + Intimate")
    numbered(doc, "Whether your APIs are IP-restricted (yes/no) — if yes, request PLRA UAT egress IP(s) for your firewall")
    numbered(doc, "Technical contact for support / failed intimations")

    add_para(doc, "PLRA returns:", size=11, bold=True, space_before=8)
    add_table(
        doc,
        ["Item", "Purpose"],
        [
            ["Prefix code", "Challan routing (e.g. 003)"],
            ["Confirmation your UAT host is on Gateway outbound allowlist", "Side A"],
            ["PLRA Gateway UAT egress IP(s)", "So you can whitelist Side B"],
            ["Joint UAT test plan", "Connectivity + unpaid/paid/intimate"],
            ["Later: Prod egress IPs + prod host allowlist", "After UAT sign-off"],
        ],
        col_widths=[8.5, 8.2],
    )

    # ===================== Appendix B =====================
    heading(doc, "21. Appendix B — Document Control", 1)
    add_table(
        doc,
        ["Ver.", "Date", "Author", "Summary of Change", "Approved By"],
        [
            ["1.0", "10-Aug-2026", "PLRA PGW Team", "Initial Department / Service Integration Guide", "Pending"],
            ["1.1", "20-Aug-2026", "PLRA PGW Team", "Formalised as Onboarding SOP with sample payloads, dual-side whitelisting, UAT and Live checklists", "Pending"],
        ],
    )

    # ===================== Appendix C =====================
    heading(doc, "22. Appendix C — Acceptance and Sign-Off", 1)
    add_para(
        doc,
        "By signing below, the parties confirm that this document constitutes the agreed onboarding "
        "SOP and technical contract for integrating the named client service with the PLRA Central "
        "Payment Gateway, and that they accept the obligations assigned herein.",
    )
    add_table(
        doc,
        ["", "PLRA (Payment Gateway)", "Client Department / Service"],
        [
            ["Organisation", "Punjab Land Records Authority (PLRA)", "<Client Organisation>"],
            ["Name", "", ""],
            ["Designation", "", ""],
            ["Signature", "", ""],
            ["Date", "", ""],
            ["Official Seal", "", ""],
        ],
        col_widths=[3.5, 6.6, 6.6],
    )

    add_para(
        doc,
        "— End of Document — PLRA Central Payment Gateway New Client / Service Onboarding Guide v1.1 —",
        size=10, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16,
    )

    out = "/workspace/PLRA_Payment_Gateway_New_Client_Onboarding_Guide.docx"
    doc.save(out)
    print(f"Saved {out}")
    return out


if __name__ == "__main__":
    build()
