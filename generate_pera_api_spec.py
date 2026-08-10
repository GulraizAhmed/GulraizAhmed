#!/usr/bin/env python3
"""
Generate PLRA → PERA DC Valuation Rate API Specification Document
following the Authority G2G Interface Specification template structure.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


GREEN = RGBColor(0x00, 0x68, 0x37)
NAVY = RGBColor(0x1A, 0x2B, 0x4A)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREEN = "E8F5EE"
HEADER_GREEN = "006837"
TABLE_HDR = "D9EAD3"


def set_run(run, size=11, bold=False, italic=False, color=BLACK, font="Calibri", mono=False):
    run.font.name = "Consolas" if mono else font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, italic=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, mono=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color, mono=mono)
    return p


def add_runs(doc, parts, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    """parts: list of (text, size, bold, italic, color[, mono])"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for part in parts:
        text, size, bold, italic, color = part[:5]
        mono = part[5] if len(part) > 5 else False
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic, color=color, mono=mono)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = GREEN if level == 1 else NAVY
        run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text, size=11, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run(r, size=size, bold=True)
        r = p.add_run(text)
        set_run(r, size=size)
    else:
        r = p.add_run(text)
        set_run(r, size=size)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=9, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=9, color=BLACK)
        shade_cell(cell, TABLE_HDR)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], str(val), size=9)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def page_break(doc):
    doc.add_page_break()


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "006837")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run("PLRA-DCVAL-API-SPEC-001  |  Version 0.2  |  Restricted  |  PLRA → PERA DC Valuation Rate Interface")
        set_run(run, size=8, color=GRAY)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Punjab Land Records Authority (PLRA)  |  Govt of the Punjab  |  Page ")
        set_run(run, size=8, color=GRAY)
        # page number field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run2 = fp.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instr)
        run2._r.append(fldChar2)
        set_run(run2, size=8, color=GRAY)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    set_run(run, size=8.5, mono=True, color=NAVY)


def kv_table(doc, pairs):
    """Two-column key/value table."""
    add_table(doc, ["Attribute", "Value"], [[k, v] for k, v in pairs], col_widths=[5.5, 11.5])


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    set_header_footer(doc)

    # ===================== COVER =====================
    add_para(doc, "API SPECIFICATION DOCUMENT", size=11, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(
        doc,
        "Interface Specification for Government-to-Government (G2G) System Integration",
        size=12, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )
    add_para(
        doc,
        "PLRA DC Valuation Rate Services for PERA Activity Register",
        size=18, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
    )
    add_hr(doc)

    cover_rows = [
        ("Document Title", "PLRA DC Valuation Rate API Specification for PERA Integration"),
        ("Document Identifier", "PLRA-DCVAL-API-SPEC-001"),
        ("Version", "0.2"),
        ("Document Status", "Draft"),
        ("Security Classification", "Restricted"),
        ("Date of Issue", "10-Aug-2026"),
        ("Service Provider", "Punjab Land Records Authority (PLRA) — DC Valuation / CLRMIS Services"),
        ("Service Consumer", "PERA — Activity Register Application"),
        ("Prepared By", "Integration / API Team, PLRA"),
        ("Reviewed By", "To be assigned — Architecture / Security / QA"),
        ("Approved By", "To be assigned — Director IT / Competent Authority, PLRA"),
        ("Document Owner", "Director IT / API Governance Custodian, PLRA"),
    ]
    kv_table(doc, cover_rows)

    add_para(
        doc,
        "This document is the property of the Punjab Land Records Authority (PLRA), Government of the Punjab. "
        "It is issued for the sole purpose of enabling PERA to design, develop, test and operate an interface "
        "to PLRA’s DC Valuation Rate services. It shall not be reproduced, redistributed or disclosed to any "
        "third party, in whole or in part, without prior written authorisation of the Document Owner.",
        size=9, italic=True, color=GRAY, space_before=8,
    )

    page_break(doc)

    # ===================== SECTION 0 =====================
    heading(doc, "0. Document Control", 1)
    heading(doc, "0.1 Document Identification", 2)
    kv_table(doc, [
        ("Document Identifier", "PLRA-DCVAL-API-SPEC-001"),
        ("Current Version", "0.2"),
        ("Supersedes", "PLRA-DCVAL-API-SPEC-001 Version 0.1"),
        ("Configuration Item Reference", "CI-PLRA-DCVAL-API-SPEC"),
        ("Repository Location", "PLRA Controlled Document Repository / Integration Library"),
        ("Retention Period", "Per PLRA records retention schedule (minimum 7 years from withdrawal)"),
        ("Review Cycle", "Annually, or upon any change to the exposed interface contract"),
    ])

    heading(doc, "0.2 Revision History", 2)
    add_table(
        doc,
        ["Ver.", "Date", "Author", "Change Reference", "Summary of Change", "Approved By"],
        [
            ["0.1", "10-Aug-2026", "PLRA Integration Team", "INIT-001", "Initial draft for PERA Activity Register DC Valuation integration", "Pending"],
            ["0.2", "10-Aug-2026", "PLRA Integration Team", "CR-002", "Aligned contracts to Square/Kila fields; added full Property Area valuation (API-10); renamed multi-khasra GetValuationMethod to API-11; standardised valuation response fields (dcRate, structureRate, location, classification, unitOfMeasurement)", "Pending"],
        ],
    )
    add_para(
        doc,
        "Table 0-1: Revision history. Minor versions denote clarification without contract change; "
        "major versions denote a breaking or behavioural change to the interface.",
        size=9, italic=True, color=GRAY,
    )

    heading(doc, "0.3 Contributors", 2)
    add_table(
        doc,
        ["Name", "Designation", "Organisation / Wing", "Contribution"],
        [
            ["PLRA Integration Team", "API / Integration", "PLRA", "API design, data model and operation specifications"],
            ["Enterprise Architecture", "Solution Architect", "PLRA", "Architecture, patterns and integration decisions"],
            ["Information Security", "Security Officer", "PLRA", "Security control design (pending formal review)"],
            ["Quality Assurance", "QA Manager", "PLRA", "Verification basis and certification criteria (pending)"],
            ["PERA Integration Focal", "Consumer Technical Lead", "PERA", "Consumer requirements and Activity Register use cases"],
        ],
    )

    heading(doc, "0.4 Review and Approval Record", 2)
    add_para(
        doc,
        "No version of this document may be released externally until every mandatory reviewer named below "
        "has recorded a disposition.",
        size=10,
    )
    add_table(
        doc,
        ["Review Stage", "Reviewing Authority", "Mandatory", "Disposition", "Date", "Signature"],
        [
            ["Architecture Review", "Chief / Enterprise Architect, PLRA", "Yes", "Pending", "—", ""],
            ["Technical Review", "System Owner / Dev Lead, PLRA", "Yes", "Pending", "—", ""],
            ["Information Security Review", "CISO / Security Officer, PLRA", "Yes", "Pending", "—", ""],
            ["Data Governance Review", "Data Protection Focal Person, PLRA", "Yes", "Pending", "—", ""],
            ["Quality Assurance Review", "QA Manager, PLRA", "Yes", "Pending", "—", ""],
            ["Final Approval", "Director IT / Competent Authority, PLRA", "Yes", "Pending", "—", ""],
        ],
    )

    heading(doc, "0.5 Distribution List", 2)
    add_table(
        doc,
        ["Recipient", "Organisation", "Copy Type", "Classification Handling"],
        [
            ["Document Owner / Director IT", "PLRA", "Controlled", "Restricted — internal custody"],
            ["PERA Integration Focal Person", "PERA", "Controlled", "Restricted — named recipient only"],
            ["PLRA API Operations", "PLRA", "Controlled", "Restricted"],
            ["PLRA Information Security", "PLRA", "Controlled", "Restricted"],
        ],
    )

    heading(doc, "0.6 Issue Readiness Confirmation", 2)
    add_table(
        doc,
        ["#", "Confirmation", "Status", "Confirmed By"],
        [
            ["1", "All placeholders populated or marked Not Applicable", "In Progress", "Document Author"],
            ["2", "Author Guidance callouts and template usage page removed", "Yes", "Document Author"],
            ["3", "No live credentials, keys, certificates or production endpoints embedded", "Yes", "Document Author"],
            ["4", "Appendices attached and version-matched to the body", "Yes", "Document Author"],
            ["5", "Security classification applied to header, footer and cover", "Yes", "Document Author"],
            ["6", "Document registered against its configuration item", "Pending", "Document Owner"],
        ],
    )

    page_break(doc)

    # ===================== TOC placeholder =====================
    heading(doc, "Table of Contents", 1)
    add_para(
        doc,
        "In Microsoft Word: place the cursor below and use References → Table of Contents → Automatic Table, "
        "or update an existing TOC field with F9 / Update entire table.",
        size=10, italic=True, color=GRAY,
    )
    for item in [
        "0. Document Control",
        "1. Introduction",
        "2. Integration Overview",
        "3. Integration Architecture",
        "4. Interface Design Principles and Conventions",
        "5. Security Architecture and Controls",
        "6. API Catalogue and Operation Specifications",
        "7. Data Model and Data Dictionary",
        "8. Transaction and State Management",
        "9. Fault Management and Error Handling",
        "10. Non-Functional Requirements",
        "11. Service Levels and Operational Support",
        "12. Verification, Testing and Certification",
        "13. Consumer Implementation and Onboarding",
        "14. Lifecycle and Maintenance",
        "15. Risk Register",
        "Appendix A — Machine-Readable Interface Definition",
        "Appendix B — Sample Payload Library",
        "Appendix C — Client Collection and Test Assets",
        "Appendix D — Consolidated Error Code Register",
        "Appendix E — Requirements Traceability Matrix",
        "Appendix F — Review Checklist",
        "Appendix G — Review Observation Log",
        "Appendix H — Acceptance and Sign-Off",
    ]:
        add_para(doc, item, size=11, space_after=2)

    page_break(doc)

    # ===================== SECTION 1 =====================
    heading(doc, "1. Introduction", 1)
    heading(doc, "1.1 Purpose", 2)
    add_para(
        doc,
        "This document defines the application programming interfaces exposed by the Punjab Land Records "
        "Authority (PLRA) DC Valuation Rate services to the PERA Activity Register application, and establishes "
        "the technical contract governing their consumption. It is the single authoritative reference for "
        "interface behaviour: where any other artefact, correspondence or verbal understanding conflicts with "
        "this document, this document prevails.",
    )
    add_para(
        doc,
        "The document serves four distinct readerships simultaneously — it is the design record for the "
        "architecture function, the build specification for PERA’s development team, the control evidence for "
        "the information security function, and the acceptance basis for quality assurance.",
    )

    heading(doc, "1.2 Scope", 2)
    heading(doc, "1.2.1 In Scope", 3)
    for t in [
        "Business capabilities for retrieving District, Tehsil, Mauza / Property Area, Khasra and DC Valuation Rate data for Rural and Urban land, as enumerated in Section 2.2.",
        "Interface architecture, transport, connectivity and environment topology.",
        "Authentication, authorisation, transport protection and message-integrity controls.",
        "Operation-level request and response contracts, field-level data definitions and reference code sets.",
        "Transaction semantics for read-oriented valuation lookup, including correlation and optional multi-Khasra valuation.",
        "Fault taxonomy, error register and consumer-side recovery obligations.",
        "Service levels, operational support model and change governance.",
        "Certification, onboarding and production authorisation procedure.",
    ]:
        bullet(doc, t)

    heading(doc, "1.2.2 Out of Scope", 3)
    for t in [
        "Internal processing logic, database design and system-internal workflows of PLRA CLRMIS / DC Valuation engines.",
        "Business processes executed wholly within the PERA Activity Register (case workflow, activity registration UI beyond API consumption).",
        "Commercial, financial or legal instruments governing the engagement, which are addressed in the referenced data sharing agreement / MoU.",
        "Mutation, Fard issuance, payment gateway, GPC, or any PLRA capability not explicitly enumerated in Section 2.2.",
        "Write / update operations against land records or valuation schedules.",
    ]:
        bullet(doc, t)

    heading(doc, "1.2.3 Boundary Statement", 3)
    add_para(
        doc,
        "The interface boundary is drawn at the PLRA API gateway ingress. PLRA is accountable for all behaviour "
        "inside that boundary; PERA is accountable for all behaviour outside it, including Activity Register UX, "
        "session management and presentation of valuation results. Responsibility for the network path between "
        "the two parties is apportioned in Section 3.5.",
    )

    heading(doc, "1.3 Intended Audience and Reading Guidance", 2)
    add_table(
        doc,
        ["Audience", "Primary Interest", "Recommended Reading Path"],
        [
            ["Business and Programme Owners", "Capability coverage and service commitments", "Sections 1, 2, 11"],
            ["Enterprise and Solution Architects", "Architecture, patterns and integration decisions", "Sections 2, 3, 4, 8"],
            ["PERA Development Teams", "Contract, data model and implementation obligations", "Sections 4–9, 13, Appendices A–C"],
            ["Information Security and Audit", "Control design and evidence", "Sections 5, 7.4, 10.5, 15"],
            ["Quality Assurance and Test", "Verification basis and acceptance criteria", "Sections 6, 9, 12"],
            ["Operations and Service Desk", "Runtime behaviour, monitoring and escalation", "Sections 10, 11, 14"],
        ],
    )

    heading(doc, "1.4 Business Context", 2)
    add_para(
        doc,
        "PERA maintains an Activity Register that requires reliable, authoritative District Collector (DC) "
        "valuation rates for land parcels across Punjab. Historically, obtaining DC rates involved fragmented "
        "manual lookups and office coordination, introducing delay and inconsistency into PERA activity workflows.",
    )
    add_para(
        doc,
        "Through this integration, PERA users select geographic and property attributes (District → Tehsil → "
        "Urban/Rural path → Mauza or Property Area → Khasra) and receive DC valuation rates directly from PLRA. "
        "For Rural scenarios, multi-Khasra selection is supported so that valuation rates for multiple Khasras "
        "are returned in a single API response. The expected outcome is accurate, timely valuation data inside "
        "the PERA Activity Register without middlemen or offline rate sheets, under a formal G2G data-sharing "
        "arrangement between PLRA and PERA.",
    )

    heading(doc, "1.5 Definitions, Acronyms and Abbreviations", 2)
    add_table(
        doc,
        ["Term / Acronym", "Expansion", "Definition in the Context of This Document"],
        [
            ["API", "Application Programming Interface", "A network-addressable contract through which PERA invokes a PLRA capability."],
            ["Provider", "—", "PLRA — the organisation exposing the interface and accountable for its behaviour."],
            ["Consumer", "—", "PERA — the organisation invoking the interface under the terms of this document."],
            ["G2G", "Government-to-Government", "An integration in which both parties are public sector entities."],
            ["PLRA", "Punjab Land Records Authority", "Service Provider; custodian of land records and DC valuation data."],
            ["PERA", "PERA", "Service Consumer; operates the Activity Register requiring DC valuation rates."],
            ["CLRMIS", "Computerized Land Records Management Information System", "PLRA system of record contributing district/tehsil/mauza/khasra and valuation data."],
            ["DC Rate", "District Collector Valuation Rate", "Official land valuation rate applicable to a Khasra / property area context."],
            ["Mauza / Mouza", "Revenue estate / village unit", "Geographic unit under a Tehsil used in the Urban path to list Khasras."],
            ["Khasra", "Survey / parcel number", "Land parcel identifier against which DC valuation is returned."],
            ["Property Area", "Classified land area unit", "Area selection under Rural path (after property classification) used to list Khasras."],
            ["JWT", "JSON Web Token", "Signed token conveying authenticated consumer identity and scopes."],
            ["mTLS", "Mutual Transport Layer Security", "TLS handshake in which both parties present and validate certificates."],
            ["SLA", "Service Level Agreement", "Committed and measurable service performance thresholds."],
            ["UAT", "User Acceptance Testing", "Consumer-executed verification preceding production authorisation."],
            ["ICD", "Interface Control Document", "This document establishing the technical contract between PLRA and PERA."],
        ],
    )

    heading(doc, "1.6 Referenced Documents", 2)
    add_table(
        doc,
        ["Ref.", "Document Title", "Identifier", "Version", "Custodian"],
        [
            ["R1", "PERA Activity Register — Business Requirements for DC Valuation Lookup", "PERA-BRS-DCVAL-001", "TBD", "PERA"],
            ["R2", "PLRA–PERA Data Sharing Agreement / Memorandum of Understanding", "PLRA-PERA-DSA-001", "TBD", "PLRA / PERA"],
            ["R3", "PLRA Enterprise Information Security Policy", "PLRA-ISP", "Current", "PLRA CISO"],
            ["R4", "PLRA API Design and Governance Standard", "PLRA-API-STD", "Current", "PLRA Architecture"],
            ["R5", "Applicable Punjab land revenue / valuation instruments", "—", "Current", "Govt of the Punjab"],
            ["R6", "OpenAPI definition (Appendix A)", "PLRA-DCVAL-OAS-001", "0.1", "PLRA API Owner"],
        ],
    )

    heading(doc, "1.7 Document Conventions", 2)
    add_table(
        doc,
        ["Convention", "Meaning"],
        [
            ["SHALL / SHALL NOT", "An absolute requirement. Non-conformance is a defect and blocks certification."],
            ["SHOULD / SHOULD NOT", "A recommendation. Departure is permitted only with recorded justification."],
            ["MAY", "An optional behaviour carrying no conformance obligation."],
            ["monospaced text", "A literal value, identifier, path fragment, header name or payload element."],
            ["Reserved (RSV)", "A field defined in the contract but not yet activated; consumers shall tolerate its presence."],
        ],
    )
    add_para(
        doc,
        "Timestamps in this document and in all payloads follow ISO 8601 with an explicit offset. "
        "Pakistan Standard Time is expressed as UTC+05:00.",
        size=10,
    )

    page_break(doc)

    # ===================== SECTION 2 =====================
    heading(doc, "2. Integration Overview", 1)
    heading(doc, "2.1 Participating Entities and Roles", 2)
    add_table(
        doc,
        ["Entity", "System", "Role", "Accountable Function", "Focal Person"],
        [
            ["Punjab Land Records Authority (PLRA)", "DC Valuation / CLRMIS API Services", "Service Provider", "Directorate IT / Land Records Systems", "To be nominated"],
            ["PERA", "Activity Register Application", "Service Consumer", "PERA IT / Integration Wing", "To be nominated"],
            ["PLRA API Gateway", "Edge Gateway", "Gateway / Security Enforcement", "PLRA API Operations", "To be nominated"],
        ],
    )

    heading(doc, "2.2 Business Capabilities Exposed", 2)
    add_para(
        doc,
        "Each capability below is realised by one or more operations specified in Section 6. "
        "The capability register is the definitive statement of scope.",
        size=10,
    )
    add_table(
        doc,
        ["Cap. ID", "Business Capability", "Description", "Realising Operation(s)", "Criticality"],
        [
            ["CAP-01", "Property Type Selection", "Expose Rural / Urban property type options for PERA Activity Register intake.", "API-01", "High"],
            ["CAP-02", "Administrative Geography Lookup", "Provide District and Tehsil lists for cascading selection.", "API-02, API-03", "High"],
            ["CAP-03", "Urban Parcel Navigation", "Provide Mauza and Khasra lists under a selected Tehsil for Urban land.", "API-04, API-05", "High"],
            ["CAP-04", "Rural Parcel Navigation", "Provide property classification, property area and Khasra lists for Rural land.", "API-06, API-07, API-08", "High"],
            ["CAP-05", "DC Valuation Rate Retrieval", "Return DC valuation rate(s) for selected Khasra(s) or full Rural Property Area, including multi-Khasra responses (GetValuationMethod).", "API-09, API-10, API-11", "High"],
        ],
    )

    heading(doc, "2.2.1 End-to-End Consumer Journey", 3)
    add_para(doc, "Common path (all journeys):", size=11, bold=True)
    bullet(doc, "Call Get Property Types → user selects Rural or Urban.")
    bullet(doc, "Call Get All Districts → user selects District.")
    bullet(doc, "Call Get Tehsils by District → user selects Tehsil.")

    add_para(doc, "Urban path:", size=11, bold=True, space_before=6)
    bullet(doc, "Call Get Mauzas by Tehsil → user selects Mauza.")
    bullet(doc, "Call Get Khasra Nos by Mauza → list returns khasraId, khasraNumber, squareId, squareNumber, kilaId, kilaNumber; user selects a Khasra.")
    bullet(doc, "Call Get Valuation by Khasra No (GetValuationMethod) with khasraId, squareId, kilaId → PLRA returns dcRate, structureRate, location, classification, unitOfMeasurement.")

    add_para(doc, "Rural path:", size=11, bold=True, space_before=6)
    bullet(doc, "User selects Property Classification (Residential, Commercial, Agriculture, Industrial) via Get Property Classifications.")
    bullet(doc, "Call Get Property Areas by Tehsil (filtered by classification) → user selects Property Area.")
    bullet(doc, "Option A — Full Property Area: Call Get Valuation by Property Area → PLRA returns dcRate, structureRate, location, classification, unitOfMeasurement.")
    bullet(doc, "Option B — Khasra/Kila selection: Call Get Khasra / Kila Nos by Property Area → user may select one or multiple Khasras (and related Kilas).")
    bullet(doc, "Then Call Get Valuation by Property Area & Khasra No(s) (GetValuationMethod) → PLRA returns DC valuation details for each selected Khasra in one response.")

    heading(doc, "2.3 Integration Objectives and Success Measures", 2)
    add_table(
        doc,
        ["Obj. ID", "Objective", "Measure", "Target", "Measurement Method"],
        [
            ["OBJ-01", "Eliminate offline DC rate lookup for PERA activities", "% of Activity Register valuation lookups completed via API", "≥ 95% within 6 months of go-live", "PERA transaction analytics"],
            ["OBJ-02", "Provide authoritative PLRA rates", "Mismatch rate vs PLRA source after reconciliation sample", "< 0.1%", "Monthly joint sample"],
            ["OBJ-03", "Support multi-Khasra Rural valuation", "Successful multi-Khasra valuation responses", "100% of valid multi-select requests", "API success metrics"],
            ["OBJ-04", "Meet response-time commitments", "95th percentile latency for read APIs", "≤ 2000 ms", "Gateway metrics"],
        ],
    )

    heading(doc, "2.4 Assumptions", 2)
    add_table(
        doc,
        ["ID", "Assumption", "Impact if Invalid", "Owner"],
        [
            ["A-01", "PERA will consume APIs only for Activity Register purposes stated in the DSA.", "Purpose limitation breach; access may be suspended.", "PERA / PLRA Data Governance"],
            ["A-02", "District / Tehsil / Mauza / Khasra master data in PLRA is authoritative and current.", "Incorrect rates or empty lists presented to PERA users.", "PLRA Land Records"],
            ["A-03", "PERA will cascade calls in the prescribed order (District → Tehsil → path-specific).", "Invalid parameter combinations and elevated 4xx rates.", "PERA Dev"],
            ["A-04", "Network whitelisting and credentials will be available before UAT.", "Certification delayed.", "Both"],
        ],
    )

    heading(doc, "2.5 Constraints", 2)
    add_table(
        doc,
        ["ID", "Constraint", "Type", "Consequence for Design"],
        [
            ["C-01", "Read-only interface; no write to land records or valuation schedules.", "Legal / Technical", "Only GET/POST-query style read operations are exposed."],
            ["C-02", "Personal / land data handling per DSA and PLRA security policy.", "Legal", "Scopes, logging minimisation and purpose limitation apply."],
            ["C-03", "Production access only after certification (Section 12).", "Operational", "Sandbox-first onboarding mandatory."],
            ["C-04", "Urdu and English labels may appear in master data; UTF-8 required.", "Technical", "Consumers shall not transliterate or corrupt non-Latin script."],
        ],
    )

    heading(doc, "2.6 Dependencies", 2)
    add_table(
        doc,
        ["ID", "Dependency", "Depends Upon", "Required By", "Status"],
        [
            ["D-01", "Signed data sharing instrument", "PLRA & PERA competent authorities", "Before sandbox credentials", "Open"],
            ["D-02", "DC Valuation data availability in PLRA source systems", "PLRA Land Records / Valuation custodians", "UAT start", "Open"],
            ["D-03", "API gateway and identity platform readiness", "PLRA API Operations", "UAT start", "Open"],
            ["D-04", "PERA Activity Register UI ready for cascading dropdowns", "PERA Dev", "UAT", "Open"],
        ],
    )

    heading(doc, "2.7 Governance and Responsibility Assignment", 2)
    add_table(
        doc,
        ["Activity", "Provider Architect", "Provider Ops", "Consumer Dev", "Security", "Steering"],
        [
            ["Interface contract definition", "A", "C", "C", "C", "I"],
            ["Credential issuance and rotation", "C", "A", "R", "A", "I"],
            ["Certification and production authorisation", "C", "R", "R", "A", "A"],
            ["Incident triage and resolution", "C", "A", "R", "C", "I"],
            ["Change request adjudication", "A", "C", "C", "C", "A"],
            ["Version deprecation", "A", "R", "I", "C", "A"],
        ],
    )
    add_para(doc, "Table 2-1: R = Responsible, A = Accountable, C = Consulted, I = Informed.", size=9, italic=True, color=GRAY)

    page_break(doc)

    # ===================== SECTION 3 =====================
    heading(doc, "3. Integration Architecture", 1)
    heading(doc, "3.1 Architectural Overview", 2)
    add_para(
        doc,
        "Logical request path: PERA Activity Register → PERA egress (whitelisted IP) → network path → "
        "PLRA edge (load balancer) → PLRA API Gateway (TLS termination, authN/authZ, rate limit, schema validation, "
        "correlation injection) → DC Valuation / CLRMIS read services → response envelope returned to PERA. "
        "Trust boundary is at the PLRA API gateway ingress. No direct database connectivity from PERA is permitted.",
    )
    add_para(doc, "Figure 3-1: Logical integration architecture (narrative).", size=9, italic=True, color=GRAY)
    code_block(
        doc,
        "PERA Activity Register\n"
        "        |  HTTPS + Bearer JWT (+ mTLS if mandated)\n"
        "        v\n"
        "PLRA API Gateway  -->  Auth / Scope / Throttle / Schema\n"
        "        |\n"
        "        v\n"
        "DC Valuation Read Services (CLRMIS-backed)\n"
        "        |\n"
        "        v\n"
        "Standard JSON response envelope → PERA",
    )

    heading(doc, "3.2 Integration Pattern and Rationale", 2)
    add_table(
        doc,
        ["Attribute", "Selection", "Rationale"],
        [
            ["Interaction Style", "Request–Response", "Synchronous cascading lookups suit dropdown-driven Activity Register UX."],
            ["Invocation Mode", "Synchronous", "User waits for District/Tehsil/Mauza/Khasra/Rate lists during form completion."],
            ["Transport Protocol", "HTTPS over TLS 1.2+", "Mandatory for G2G data in transit."],
            ["Architectural Style", "REST", "Simple resource reads; aligns with PLRA API standard."],
            ["Payload Format", "JSON", "Lightweight for mobile/web Activity Register clients."],
            ["Mediation", "API Gateway", "Centralised security, metering, versioning and threat protection."],
            ["Callback Mechanism", "Not applicable", "No asynchronous callbacks required for valuation lookup."],
            ["Alternatives Rejected", "SOAP / file extract / DB link", "Heavier coupling; DB link violates boundary and security policy."],
        ],
    )

    heading(doc, "3.3 Environment Topology", 2)
    add_table(
        doc,
        ["Environment", "Purpose", "Base URL", "Data Character", "Availability Window", "Access Control"],
        [
            ["Development", "Provider-internal build", "https://dev-api.plra.punjab.gov.pk (illustrative)", "Synthetic", "Provider business hours", "Provider only"],
            ["Sandbox / UAT", "Consumer integration and certification", "https://uat-api.plra.punjab.gov.pk (illustrative)", "Synthetic / masked", "Agreed UAT window", "Whitelisted PERA"],
            ["Staging", "Pre-production rehearsal", "https://stg-api.plra.punjab.gov.pk (illustrative)", "Masked", "Restricted", "Restricted"],
            ["Production", "Live service", "https://api.plra.punjab.gov.pk (illustrative)", "Live", "24×7 less maintenance", "Whitelisted PERA"],
        ],
    )
    add_para(
        doc,
        "Illustrative hostnames only. Actual endpoints are issued out of band with credentials. "
        "Production data SHALL NOT be present in non-production environments.",
        size=9, italic=True, color=GRAY,
    )

    heading(doc, "3.4 Network Connectivity", 2)
    add_table(
        doc,
        ["Parameter", "Specification"],
        [
            ["Connectivity Model", "Whitelisted public internet over HTTPS, or site-to-site IPSec VPN if mandated by security review"],
            ["Provider Ingress Endpoints", "FQDN and resolved IPs per environment — issued out of band"],
            ["Consumer Egress Addresses", "PERA static source IP addresses / CIDR ranges to be whitelisted"],
            ["Permitted Ports and Protocols", "TCP 443, HTTPS only"],
            ["Firewall Change Procedure", "PERA submits connectivity request; PLRA network function raises change; both confirm reachability"],
            ["Bandwidth Provisioning", "Best-effort over internet; dedicated circuit if jointly agreed"],
            ["Redundancy", "Provider multi-AZ / secondary path per PLRA operations standard"],
            ["DNS Resolution", "Public DNS for published FQDNs unless private zone agreed"],
        ],
    )

    heading(doc, "3.4.1 Connectivity Establishment Prerequisites", 3)
    for i, t in enumerate([
        "PERA submits static egress addresses and the signed connectivity request to the PLRA network function.",
        "PLRA raises the firewall change under its change management procedure and confirms the rule reference.",
        "Both parties execute a bidirectional reachability test and record the result.",
        "Certificates / client credentials are exchanged through the channel defined in Section 5.6; connectivity is not deemed established until a successful authenticated call is logged.",
    ], 1):
        add_para(doc, f"{i}. {t}", size=10, space_after=3)

    heading(doc, "3.5 Boundary Responsibility Demarcation", 2)
    add_table(
        doc,
        ["Segment", "Owning Party", "Monitoring Responsibility", "Fault Ownership"],
        [
            ["PERA application and egress", "PERA", "PERA", "PERA"],
            ["Intervening network / carrier", "Carrier / Shared", "Both", "Escalation path per Section 11"],
            ["PLRA ingress and gateway", "PLRA", "PLRA", "PLRA"],
            ["PLRA application and datastore", "PLRA", "PLRA", "PLRA"],
        ],
    )

    heading(doc, "3.6 Gateway Responsibilities", 2)
    for t in [
        "Termination of transport security and enforcement of the minimum protocol version and cipher policy.",
        "Authentication of the caller and validation of token integrity, expiry, issuer and audience.",
        "Authorisation against the scope model defined in Section 5.3.",
        "Traffic management: rate limiting, quota enforcement, concurrency capping and payload size limits.",
        "Threat protection: schema validation, injection screening and replay detection where applicable.",
        "Correlation identifier injection/echo, request logging and metering.",
        "Routing, version resolution and response caching where declared per operation.",
    ]:
        bullet(doc, t)

    heading(doc, "3.7 Architectural Decision Record", 2)
    add_table(
        doc,
        ["ADR ID", "Decision", "Status", "Rationale", "Consequence"],
        [
            ["ADR-01", "REST/JSON over HTTPS via API Gateway", "Accepted", "Fits cascading read UX; aligns with PLRA API standard", "Gateway-enforced security and versioning"],
            ["ADR-02", "Separate Urban (Mauza) and Rural (Property Area) navigation paths", "Accepted", "Matches PERA Activity Register business rules", "Two valuation operations (single vs multi-Khasra)"],
            ["ADR-03", "Multi-Khasra valuation supported only on Rural path via GetValuationMethod", "Accepted", "PERA Rural use case requires multi-select rates", "API-11 accepts khasras[] with optional squareId/kilaId; API-10 covers full-area valuation"],
            ["ADR-04", "Read-only scopes; no mutation APIs", "Accepted", "Purpose limited to rate lookup", "Reduced attack surface"],
        ],
    )

    page_break(doc)

    # ===================== SECTION 4 =====================
    heading(doc, "4. Interface Design Principles and Conventions", 1)
    heading(doc, "4.1 Design Tenets", 2)
    add_table(
        doc,
        ["Tenet", "Statement"],
        [
            ["Contract stability", "A published contract is immutable. Change is delivered through versioning."],
            ["Consumer independence", "The contract exposes business capability, not PLRA internal tables."],
            ["Explicitness", "Defaults, optionality and null semantics are declared for every field."],
            ["Predictable failure", "Every failure mode returns a deterministic, documented response."],
            ["Least privilege", "PERA is granted only DC valuation read scopes."],
            ["Auditability", "Every invocation is attributable to an authenticated identity."],
        ],
    )

    heading(doc, "4.2 Resource and Endpoint Naming", 2)
    for t in [
        "Resources are named as plural nouns where applicable; controller-style actions use clear verb nouns only when required.",
        "Path segments are lower-case, hyphen-delimited; query parameters and payload fields are camelCase.",
        "Path structure: /dc-valuation/v1/{resource}[/{identifier}][/{sub-resource}]",
        "Actual paths are specified per operation in Section 6.",
    ]:
        bullet(doc, t)

    heading(doc, "4.3 HTTP Method Semantics", 2)
    add_table(
        doc,
        ["Method", "Applied To", "Safe", "Idempotent", "Typical Success Status"],
        [
            ["GET", "Retrieval of master lists and single valuation lookups", "Yes", "Yes", "200"],
            ["POST", "Valuation lookup with khasra context body (API-09, API-11)", "Yes*", "Yes*", "200"],
        ],
    )
    add_para(
        doc,
        "* API-09 and API-11 are treated as safe, idempotent queries that use POST solely to carry "
        "structured khasra context in the request body. They SHALL NOT mutate Provider state.",
        size=9, italic=True, color=GRAY,
    )

    heading(doc, "4.4 Request Correlation", 2)
    add_para(
        doc,
        "Every request SHALL carry a consumer-generated X-Correlation-Id (UUID v4). PLRA echoes it on every "
        "response and records it in the audit trail. PERA SHALL retain the correlation identifier against the "
        "Activity Register record for the period stated in Section 10.6.",
    )

    heading(doc, "4.5 Collection Handling", 2)
    add_table(
        doc,
        ["Concern", "Convention", "Default", "Bound"],
        [
            ["Pagination", "Offset-based via page & size query params where collections are large", "page=1, size=50", "Maximum size=200"],
            ["Sorting", "Provider default sort (name/code ascending)", "name asc", "As documented per operation"],
            ["Filtering", "Path/query parameters (districtId, tehsilId, etc.)", "None", "Only documented filters"],
            ["Sparse Field Selection", "Not Supported in v1", "All fields", "—"],
            ["Result Ceiling", "Maximum 200 records per page", "—", "Use pagination; 400 if size exceeded"],
        ],
    )

    heading(doc, "4.6 Versioning and Deprecation", 2)
    add_table(
        doc,
        ["Aspect", "Policy"],
        [
            ["Versioning Scheme", "URI path versioning — /dc-valuation/v1/"],
            ["Breaking Change", "Field removal/rename, narrowed domain, new mandatory input, semantic change, error code meaning change"],
            ["Non-Breaking Change", "Optional input, new response field, new extensible code value, new operation"],
            ["Consumer Obligation", "PERA SHALL implement tolerant reading and SHALL NOT fail on unrecognised response fields"],
            ["Concurrent Versions", "Minimum two major versions during overlap when a breaking change is released"],
            ["Deprecation Notice", "Minimum 180 calendar days before withdrawal"],
            ["Deprecation Signalling", "Deprecation and Sunset response headers plus formal written notice"],
            ["Withdrawal Procedure", "Architecture approval, consumer migration confirmation, gateway disablement"],
        ],
    )

    heading(doc, "4.7 Data Representation Conventions", 2)
    add_table(
        doc,
        ["Data Type", "Representation", "Example Format", "Notes"],
        [
            ["Date", "ISO 8601 calendar date", "YYYY-MM-DD", "No timezone component"],
            ["Timestamp", "ISO 8601 with offset", "YYYY-MM-DDThh:mm:ss+05:00", "PKT is +05:00"],
            ["Monetary / Rate Amount", "Decimal string", "precision as returned by PLRA", "Currency PKR unless stated"],
            ["Boolean", "JSON literal", "true / false", "Quoted strings rejected"],
            ["Enumeration", "Upper-case token", "URBAN / RURAL", "Governed by Section 7.3"],
            ["Identifiers", "String or integer as documented", "—", "Treat as opaque; do not invent"],
            ["Free Text / Names", "UTF-8", "—", "Urdu script supported; no transliteration"],
            ["Null Semantics", "Absent vs null", "—", "Absent = not supplied; null = explicitly empty"],
        ],
    )

    heading(doc, "4.8 Character Encoding and Localisation", 2)
    add_para(
        doc,
        "All payloads are encoded in UTF-8. Fields carrying Urdu or other non-Latin script SHALL be transmitted "
        "without transliteration. PERA SHALL persist and render such fields in the encoding received. "
        "The Accept-Language header is not honoured for message localisation in v1; English diagnostic messages are returned.",
    )

    heading(doc, "4.9 Technology Tailoring Record", 2)
    add_para(
        doc,
        "Not Applicable — this interface is realised as REST/JSON over HTTPS. No SOAP, GraphQL, message-queue "
        "or file-based substitutions apply in Version 0.1.",
        size=10, italic=True,
    )

    page_break(doc)

    # ===================== SECTION 5 =====================
    heading(doc, "5. Security Architecture and Controls", 1)
    add_para(
        doc,
        "This section is reviewed independently by the information security function. Statements describe "
        "controls that are implemented and testable. No key material, credential, certificate or secret "
        "appears in this document.",
        size=10, italic=True, color=GRAY,
    )

    heading(doc, "5.1 Security Objectives", 2)
    add_table(
        doc,
        ["Objective", "Control Statement", "Realised By"],
        [
            ["Confidentiality", "Data in transit is protected against interception and disclosure.", "Section 5.4"],
            ["Integrity", "Messages are protected against undetected modification and replay.", "Sections 5.4, 5.5"],
            ["Authenticity", "Every request is attributable to a registered, verified PERA client identity.", "Section 5.2"],
            ["Authorisation", "PERA can invoke only granted DC valuation read capabilities.", "Section 5.3"],
            ["Non-repudiation", "Invocations are attributable via audit logs and correlation IDs.", "Sections 5.5, 10.5"],
            ["Availability", "The service is protected against exhaustion and abusive traffic.", "Section 5.7"],
        ],
    )

    heading(doc, "5.2 Authentication", 2)
    add_table(
        doc,
        ["Parameter", "Specification"],
        [
            ["Mechanism", "OAuth 2.0 Client Credentials (machine-to-machine); mTLS MAY be mandated per security review"],
            ["Identity Granularity", "One client identity per PERA system per environment"],
            ["Token Endpoint", "Issued per environment out of band (not embedded here)"],
            ["Grant Type", "client_credentials"],
            ["Token Format", "JWT (JWS), signed with RS256"],
            ["Token Lifetime", "3600 seconds (subject to IdP policy)"],
            ["Refresh Behaviour", "Re-issue on expiry; refresh tokens not issued for M2M grants"],
            ["Clock Skew Tolerance", "±120 seconds"],
            ["Token Caching Obligation", "PERA SHALL cache and reuse until expiry; per-request token requests are prohibited"],
            ["Client Authentication", "Client secret and/or private key JWT / TLS client certificate as issued"],
            ["Failure Response", "HTTP 401 with ERR-2001 / ERR-2002 per Section 9.3"],
        ],
    )

    heading(doc, "5.2.1 Token Claim Set", 3)
    add_table(
        doc,
        ["Claim", "Description", "Validation Performed by Provider"],
        [
            ["iss", "Issuing authority", "Matched against registered issuer"],
            ["sub", "PERA client identifier", "Resolved to an active registration"],
            ["aud", "Intended audience", "Matched against DC Valuation service identifier"],
            ["exp", "Expiry timestamp", "Rejected if elapsed beyond permitted skew"],
            ["iat", "Issuance timestamp", "Rejected if future-dated beyond permitted skew"],
            ["jti", "Unique token identifier", "Checked against replay cache where enabled"],
            ["scope", "Granted authorisation scopes", "Evaluated per Section 5.3"],
        ],
    )

    heading(doc, "5.3 Authorisation Model", 2)
    add_table(
        doc,
        ["Scope", "Grants", "Applies To", "Assigned To"],
        [
            ["dcval.read.masters", "Read District, Tehsil, Mauza, Property Area, Classification, Khasra lists", "API-01 to API-08", "PERA Activity Register client"],
            ["dcval.read.valuation", "Read DC valuation rate(s)", "API-09, API-10, API-11", "PERA Activity Register client"],
        ],
    )
    add_para(
        doc,
        "No write or admin scopes are granted. Jurisdictional restriction (e.g. District-limited access) MAY be "
        "applied later via DSA amendment; if applied, it will be documented in a minor/major revision as appropriate.",
        size=10,
    )

    heading(doc, "5.4 Transport Security", 2)
    add_table(
        doc,
        ["Control", "Requirement"],
        [
            ["Minimum Protocol", "TLS 1.2; TLS 1.3 preferred. Earlier versions disabled."],
            ["Cipher Policy", "Approved cipher suites; forward secrecy required"],
            ["Certificate Authority", "Issuing CA per environment — communicated out of band"],
            ["Certificate Validation", "Chain, hostname, validity, revocation checking"],
            ["Mutual TLS", "MAY be required following security review; issuance procedure in Section 5.6"],
            ["Certificate Pinning", "Not required for v1 unless mandated by security review"],
            ["Plaintext Fallback", "Prohibited. Unprotected channel requests are rejected without processing."],
        ],
    )

    heading(doc, "5.5 Message-Level Protection", 2)
    add_table(
        doc,
        ["Control", "Specification"],
        [
            ["Request Signing", "Not applicable for v1 beyond TLS (+ optional mTLS)"],
            ["Timestamp Header", "X-Request-Timestamp MAY be required; ISO 8601 with offset; max age 5 minutes if enforced"],
            ["Nonce Handling", "Not applicable for v1 read APIs unless replay controls are extended"],
            ["Payload Encryption", "Not applicable — TLS provides confidentiality"],
            ["Response Signing", "Not applicable for v1"],
        ],
    )

    heading(doc, "5.6 Credential and Key Lifecycle", 2)
    add_table(
        doc,
        ["Stage", "Procedure", "Responsible Party", "Evidence Retained"],
        [
            ["Request", "Formal credential request after DSA and registration", "PERA → PLRA Ops", "Request record"],
            ["Issuance", "Generate and transmit via secure out-of-band channel", "PLRA Ops", "Handover record"],
            ["Storage", "Secrets vault / HSM; never in source control", "PERA", "Storage attestation"],
            ["Rotation", "At least annually or on suspicion of compromise; overlap window agreed", "Both", "Rotation log"],
            ["Revocation", "On compromise, offboarding or DSA breach; target ≤ 4 hours", "PLRA Ops", "Revocation record"],
            ["Compromise Response", "Notify PLRA Security within 1 hour; suspend client; re-issue", "Both", "Incident record"],
        ],
    )

    heading(doc, "5.7 Threat Protection Controls", 2)
    add_table(
        doc,
        ["Threat", "Control", "Configured Threshold", "Response on Breach"],
        [
            ["Credential brute force", "Failed-authentication lockout", "Per IdP policy", "401 / temporary lockout"],
            ["Replay", "Token jti / short-lived tokens", "Token lifetime", "401"],
            ["Volumetric abuse", "Rate limiting and quota", "See Section 10.2", "429 with Retry-After"],
            ["Oversized payload", "Request size ceiling", "64 KB for v1", "413"],
            ["Injection", "Schema validation", "Strict schema", "400 with validation detail"],
            ["Enumeration / scraping", "Result ceiling + anomaly detection", "page size ≤ 200", "Throttle / investigate"],
            ["Denial of service", "Edge protection and concurrency cap", "Per gateway policy", "503 / drop"],
        ],
    )

    heading(doc, "5.8 Compliance, Audit and Assurance", 2)
    for t in [
        "Every request and response is logged with correlation identifier, authenticated identity, source address, operation, outcome and latency. Payload logging is restricted to non-sensitive fields / identifiers; full land-owner PII is not logged in gateway access logs.",
        "Audit records are retained per Section 10.6 and are tamper-evident.",
        "The interface is subject to vulnerability assessment / penetration testing prior to production authorisation and thereafter annually or on major change.",
        "Applicable instruments: PLRA ISP (R3), DSA/MoU (R2), and applicable Punjab data protection / land records regulations (R5).",
        "Findings are tracked to closure in the risk register at Section 15.",
    ]:
        bullet(doc, t)

    page_break(doc)

    # ===================== SECTION 6 =====================
    heading(doc, "6. API Catalogue and Operation Specifications", 1)
    heading(doc, "6.1 Operation Inventory", 2)
    add_table(
        doc,
        ["API ID", "Operation Name", "Method", "Resource Path", "Capability", "State Changing", "Scope"],
        [
            ["API-01", "Get Property Types", "GET", "/dc-valuation/v1/property-types", "CAP-01", "No", "dcval.read.masters"],
            ["API-02", "Get All Districts", "GET", "/dc-valuation/v1/districts", "CAP-02", "No", "dcval.read.masters"],
            ["API-03", "Get Tehsils by District", "GET", "/dc-valuation/v1/districts/{districtId}/tehsils", "CAP-02", "No", "dcval.read.masters"],
            ["API-04", "Get Mauzas by Tehsil", "GET", "/dc-valuation/v1/tehsils/{tehsilId}/mauzas", "CAP-03", "No", "dcval.read.masters"],
            ["API-05", "Get Khasra Nos by Mauza", "GET", "/dc-valuation/v1/mauzas/{mauzaId}/khasras", "CAP-03", "No", "dcval.read.masters"],
            ["API-06", "Get Property Classifications", "GET", "/dc-valuation/v1/property-classifications", "CAP-04", "No", "dcval.read.masters"],
            ["API-07", "Get Property Areas by Tehsil", "GET", "/dc-valuation/v1/tehsils/{tehsilId}/property-areas", "CAP-04", "No", "dcval.read.masters"],
            ["API-08", "Get Khasra / Kila Nos by Property Area", "GET", "/dc-valuation/v1/property-areas/{propertyAreaId}/khasras", "CAP-04", "No", "dcval.read.masters"],
            ["API-09", "Get Valuation by Khasra No (GetValuationMethod)", "POST", "/dc-valuation/v1/valuations/by-khasra", "CAP-05", "No", "dcval.read.valuation"],
            ["API-10", "Get Valuation by Property Area (full area)", "GET", "/dc-valuation/v1/property-areas/{propertyAreaId}/valuation", "CAP-05", "No", "dcval.read.valuation"],
            ["API-11", "Get Valuation by Property Area & Khasra No(s) (GetValuationMethod)", "POST", "/dc-valuation/v1/valuations/by-property-area", "CAP-05", "No", "dcval.read.valuation"],
        ],
    )
    add_para(
        doc,
        "Table 6-1: Operation inventory. Machine-readable definition at Appendix A SHALL agree with this table. "
        "API-09 and API-11 realise the business operation commonly referred to as GetValuationMethod "
        "(single Urban khasra context vs Rural multi-khasra / property-area context).",
        size=9, italic=True, color=GRAY,
    )

    heading(doc, "6.2 Common Request Headers", 2)
    add_table(
        doc,
        ["Header", "Obligation", "Format / Permitted Values", "Description"],
        [
            ["Authorization", "SHALL", "Bearer {access_token}", "Credential established under Section 5.2"],
            ["Content-Type", "SHALL for requests with body", "application/json; charset=utf-8", "Required for API-09 and API-11"],
            ["Accept", "SHALL", "application/json", "Content negotiation"],
            ["X-Correlation-Id", "SHALL", "UUID v4", "Consumer-generated; echoed on response"],
            ["X-Request-Timestamp", "MAY / SHALL if enforced", "ISO 8601 with offset", "Freshness validation"],
            ["X-Consumer-Id", "MAY", "Registered PERA client id", "Consumer system identity"],
        ],
    )

    heading(doc, "6.3 Common Response Headers", 2)
    add_table(
        doc,
        ["Header", "Description"],
        [
            ["X-Correlation-Id", "Echo of the value supplied by the consumer"],
            ["X-Transaction-Id", "Provider-generated identifier for the processed request"],
            ["X-RateLimit-Remaining", "Requests remaining in the current quota interval"],
            ["Retry-After", "Seconds to wait before retrying, returned with 429 and 503"],
            ["Deprecation / Sunset", "Returned when the invoked version is scheduled for withdrawal"],
        ],
    )

    heading(doc, "6.4 Standard Response Envelope", 2)
    code_block(
        doc,
        '{\n'
        '  "status": "SUCCESS | FAILURE",\n'
        '  "responseCode": "<canonical code, Section 9.3>",\n'
        '  "responseMessage": "<human-readable, non-authoritative>",\n'
        '  "correlationId": "<echo of X-Correlation-Id>",\n'
        '  "transactionId": "<provider transaction reference>",\n'
        '  "timestamp": "YYYY-MM-DDThh:mm:ss+05:00",\n'
        '  "data": { },\n'
        '  "errors": [ { "code": "<code>", "field": "<field>", "message": "<detail>" } ]\n'
        '}',
    )
    for t in [
        "responseCode is the authoritative outcome indicator and the only element on which consumer logic SHALL branch.",
        "responseMessage is diagnostic and MAY change without a version increment.",
        "errors is present only where status is FAILURE.",
        "data is null on failure.",
    ]:
        bullet(doc, t)

    # ---- Operation helper ----
    def add_operation(api_id, name, purpose, cap, method, path, scope, preconds, postconds,
                      params, body_fields, sample_req, resp_fields, sample_resp, codes, rules, notes):
        heading(doc, f"6.5 Operation Specification — {api_id} {name}", 2)
        heading(doc, "Operation Identification", 3)
        kv_table(doc, [
            ("API Identifier", api_id),
            ("Operation Name", name),
            ("Business Purpose", purpose),
            ("Traces To Capability", cap),
            ("HTTP Method", method),
            ("Resource Path", path),
            ("Full Endpoint (UAT)", f"{{uat-base}}{path}"),
            ("Full Endpoint (Production)", f"{{prod-base}}{path}"),
            ("Required Scope", scope),
            ("State Changing", "No"),
            ("Idempotency", "Naturally idempotent"),
            ("Expected Invocation Volume", "Aligned to PERA Activity Register usage; peak during business hours"),
            ("Target Response Time", "95th percentile ≤ 2000 ms"),
            ("Timeout", "Provider processing timeout 10 seconds"),
            ("Cacheable", "Yes for master data lists (short TTL at gateway); No for live valuation if rates are volatile — see notes"),
            ("Introduced In Version", "0.1 / v1"),
            ("Status", "Active (Draft contract)"),
        ])

        heading(doc, "Pre-conditions", 3)
        for t in preconds:
            bullet(doc, t)

        heading(doc, "Post-conditions", 3)
        for t in postconds:
            bullet(doc, t)

        heading(doc, "Request Specification", 3)
        if params:
            add_table(
                doc,
                ["Parameter", "In", "Type", "Obligation", "Constraint", "Example", "Description"],
                params,
            )
        else:
            add_para(doc, "No path or query parameters.", size=10, italic=True)

        if body_fields:
            add_para(doc, "Request body fields:", size=10, bold=True)
            add_table(
                doc,
                ["Field", "Type", "Length / Format", "Obligation", "Validation Rule", "Example"],
                body_fields,
            )
        add_para(doc, "Sample request:", size=10, bold=True)
        code_block(doc, sample_req)

        heading(doc, "Response Specification", 3)
        add_table(
            doc,
            ["Field", "Type", "Length / Format", "Always Present", "Description"],
            resp_fields,
        )
        add_para(doc, "Sample success response:", size=10, bold=True)
        code_block(doc, sample_resp)

        heading(doc, "Operation-Specific Response Codes", 3)
        add_table(
            doc,
            ["HTTP Status", "Response Code", "Condition", "Consumer Action"],
            codes,
        )

        heading(doc, "Business Rules Applied", 3)
        add_table(
            doc,
            ["Rule ID", "Rule Statement", "Enforcement Point", "Violation Code"],
            rules,
        )

        heading(doc, "Sequence of Interaction", 3)
        add_para(doc, "Simple synchronous request–response. No callback or compensating action.", size=10)

        heading(doc, "Dependencies and Failure Propagation", 3)
        add_table(
            doc,
            ["Downstream Dependency", "Invoked", "Failure Behaviour", "Consumer-Visible Effect"],
            [["CLRMIS / DC Valuation read store", "Synchronously", "Fail fast", "ERR-5001 / 503 as applicable"]],
        )

        heading(doc, "Operation Notes and Known Limitations", 3)
        for t in notes:
            bullet(doc, t)

    # API-01
    add_operation(
        "API-01", "Get Property Types",
        "Returns the property type options (RURAL, URBAN) for PERA Activity Register selection.",
        "CAP-01", "GET", "/dc-valuation/v1/property-types", "dcval.read.masters",
        ["Caller is authenticated with scope dcval.read.masters."],
        ["No Provider state change. Consumer receives the property type list."],
        [],
        None,
        "GET /dc-valuation/v1/property-types HTTP/1.1\n"
        "Host: {uat-host}\n"
        "Authorization: Bearer {token}\n"
        "Accept: application/json\n"
        "X-Correlation-Id: 8f1c2e3a-4b5d-6e7f-8091-a2b3c4d5e6f7",
        [
            ["propertyTypes[].code", "string", "RURAL | URBAN", "Yes", "Property type code"],
            ["propertyTypes[].name", "string", "UTF-8", "Yes", "Display name"],
        ],
        'HTTP/1.1 200 OK\n'
        'Content-Type: application/json; charset=utf-8\n'
        'X-Correlation-Id: 8f1c2e3a-4b5d-6e7f-8091-a2b3c4d5e6f7\n\n'
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "responseMessage": "Property types retrieved successfully",\n'
        '  "correlationId": "8f1c2e3a-4b5d-6e7f-8091-a2b3c4d5e6f7",\n'
        '  "transactionId": "TXN-20260810-000001",\n'
        '  "timestamp": "2026-08-10T15:30:00+05:00",\n'
        '  "data": {\n'
        '    "propertyTypes": [\n'
        '      { "code": "URBAN", "name": "Urban" },\n'
        '      { "code": "RURAL", "name": "Rural" }\n'
        '    ]\n'
        '  },\n'
        '  "errors": null\n'
        '}',
        [
            ["200", "DCVAL-0000", "Success", "Populate Urban/Rural selection"],
            ["401", "ERR-2001", "Invalid/expired token", "Re-authenticate"],
            ["403", "ERR-3001", "Missing scope", "Escalate; do not retry"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [["BR-01", "Only RURAL and URBAN are valid property types in v1", "Service", "N/A"]],
        ["Values drive which subsequent navigation APIs PERA SHALL call."],
    )

    # API-02
    add_operation(
        "API-02", "Get All Districts",
        "Returns all Districts available for DC valuation navigation in Punjab.",
        "CAP-02", "GET", "/dc-valuation/v1/districts", "dcval.read.masters",
        ["Caller authenticated with dcval.read.masters."],
        ["No state change. District list returned for dropdown binding."],
        [["page", "query", "integer", "O", "≥1", "1", "Page number"],
         ["size", "query", "integer", "O", "1–200", "50", "Page size"]],
        None,
        "GET /dc-valuation/v1/districts?page=1&size=50 HTTP/1.1\n"
        "Host: {uat-host}\n"
        "Authorization: Bearer {token}\n"
        "Accept: application/json\n"
        "X-Correlation-Id: 1a2b3c4d-5e6f-7081-9243-54657687980a",
        [
            ["districts[].districtId", "string", "opaque id", "Yes", "District identifier"],
            ["districts[].districtCode", "string", "code", "Yes", "District code"],
            ["districts[].districtName", "string", "UTF-8", "Yes", "District name (may include Urdu)"],
            ["page", "object", "—", "Yes", "Pagination metadata (page, size, totalElements, totalPages)"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "responseMessage": "Districts retrieved successfully",\n'
        '  "correlationId": "1a2b3c4d-5e6f-7081-9243-54657687980a",\n'
        '  "transactionId": "TXN-20260810-000002",\n'
        '  "timestamp": "2026-08-10T15:31:00+05:00",\n'
        '  "data": {\n'
        '    "districts": [\n'
        '      { "districtId": "D-042", "districtCode": "LHR", "districtName": "Lahore" }\n'
        '    ],\n'
        '    "page": { "page": 1, "size": 50, "totalElements": 36, "totalPages": 1 }\n'
        '  },\n'
        '  "errors": null\n'
        '}',
        [
            ["200", "DCVAL-0000", "Success", "Bind District dropdown"],
            ["400", "ERR-1001", "Invalid page/size", "Correct and resubmit"],
            ["401", "ERR-2001", "Auth failure", "Re-authenticate"],
            ["403", "ERR-3001", "Missing scope", "Escalate"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [["BR-02", "Only active districts for valuation navigation are returned", "Service", "N/A"]],
        ["PERA SHALL use districtId from this response as input to API-03."],
    )

    # API-03
    add_operation(
        "API-03", "Get Tehsils by District",
        "Returns Tehsils belonging to the selected District for cascading selection.",
        "CAP-02", "GET", "/dc-valuation/v1/districts/{districtId}/tehsils", "dcval.read.masters",
        ["Valid districtId previously obtained from API-02.", "Caller authenticated with dcval.read.masters."],
        ["No state change. Tehsil list returned."],
        [["districtId", "path", "string", "M", "Must exist", "D-042", "District identifier"],
         ["page", "query", "integer", "O", "≥1", "1", "Page number"],
         ["size", "query", "integer", "O", "1–200", "50", "Page size"]],
        None,
        "GET /dc-valuation/v1/districts/D-042/tehsils?page=1&size=50 HTTP/1.1\n"
        "Host: {uat-host}\n"
        "Authorization: Bearer {token}\n"
        "Accept: application/json\n"
        "X-Correlation-Id: 9b8a7c6d-5e4f-3210-aabb-ccddeeff0011",
        [
            ["tehsils[].tehsilId", "string", "opaque id", "Yes", "Tehsil identifier"],
            ["tehsils[].tehsilCode", "string", "code", "Yes", "Tehsil code"],
            ["tehsils[].tehsilName", "string", "UTF-8", "Yes", "Tehsil name"],
            ["tehsils[].districtId", "string", "opaque id", "Yes", "Parent district"],
            ["page", "object", "—", "Yes", "Pagination metadata"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "data": {\n'
        '    "tehsils": [\n'
        '      { "tehsilId": "T-010", "tehsilCode": "CNT", "tehsilName": "City", "districtId": "D-042" }\n'
        '    ],\n'
        '    "page": { "page": 1, "size": 50, "totalElements": 5, "totalPages": 1 }\n'
        '  }\n'
        '}',
        [
            ["200", "DCVAL-0000", "Success", "Bind Tehsil dropdown"],
            ["404", "ERR-4002", "Unknown districtId", "Refresh districts; do not retry same id"],
            ["400", "ERR-1001", "Validation failure", "Correct request"],
            ["401", "ERR-2001", "Auth failure", "Re-authenticate"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [["BR-03", "Tehsils returned SHALL belong to the path districtId", "Service", "ERR-4002"]],
        ["After Tehsil selection, PERA branches on Urban vs Rural (API-01 selection)."],
    )

    # API-04 Urban Mauzas
    add_operation(
        "API-04", "Get Mauzas by Tehsil",
        "Returns Mauzas under the selected Tehsil for the Urban navigation path.",
        "CAP-03", "GET", "/dc-valuation/v1/tehsils/{tehsilId}/mauzas", "dcval.read.masters",
        ["Valid tehsilId from API-03.", "PERA user has selected property type URBAN.", "Caller authenticated."],
        ["No state change. Mauza list returned for Urban path."],
        [["tehsilId", "path", "string", "M", "Must exist", "T-010", "Tehsil identifier"],
         ["propertyType", "query", "string", "M", "URBAN", "URBAN", "Shall be URBAN for this operation"],
         ["page", "query", "integer", "O", "1–200 size", "1", "Page number"],
         ["size", "query", "integer", "O", "1–200", "50", "Page size"]],
        None,
        "GET /dc-valuation/v1/tehsils/T-010/mauzas?propertyType=URBAN&page=1&size=50 HTTP/1.1\n"
        "Authorization: Bearer {token}\n"
        "X-Correlation-Id: {uuid}",
        [
            ["mauzas[].mauzaId", "string", "opaque id", "Yes", "Mauza identifier"],
            ["mauzas[].mauzaCode", "string", "code", "Yes", "Mauza code"],
            ["mauzas[].mauzaName", "string", "UTF-8", "Yes", "Mauza name"],
            ["mauzas[].tehsilId", "string", "opaque id", "Yes", "Parent tehsil"],
            ["page", "object", "—", "Yes", "Pagination metadata"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "data": {\n'
        '    "mauzas": [\n'
        '      { "mauzaId": "M-100", "mauzaCode": "MZ-100", "mauzaName": "Sample Mauza", "tehsilId": "T-010" }\n'
        '    ]\n'
        '  }\n'
        '}',
        [
            ["200", "DCVAL-0000", "Success", "Bind Mauza dropdown"],
            ["400", "ERR-1002", "propertyType missing/invalid", "Send propertyType=URBAN"],
            ["404", "ERR-4002", "Unknown tehsilId", "Refresh tehsils"],
            ["422", "ERR-4001", "propertyType not URBAN", "Use Rural APIs instead"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [["BR-04", "Mauza list for Urban path requires propertyType=URBAN", "Service", "ERR-4001"]],
        ["For Rural path, PERA SHALL NOT use this operation; use API-06/API-07 instead."],
    )

    # API-05
    add_operation(
        "API-05", "Get Khasra Nos by Mauza",
        "Returns Khasra numbers (with Square and Kila identifiers) associated with the selected Mauza (Urban path).",
        "CAP-03", "GET", "/dc-valuation/v1/mauzas/{mauzaId}/khasras", "dcval.read.masters",
        ["Valid mauzaId from API-04.", "Urban path in progress."],
        ["No state change. Khasra list with Square/Kila identifiers returned."],
        [["mauzaId", "path", "string", "M", "Must exist", "M-100", "Mauza identifier"],
         ["page", "query", "integer", "O", "≥1", "1", "Page"],
         ["size", "query", "integer", "O", "1–200", "50", "Size"]],
        None,
        "GET /dc-valuation/v1/mauzas/M-100/khasras?page=1&size=50 HTTP/1.1\n"
        "Authorization: Bearer {token}\n"
        "X-Correlation-Id: {uuid}",
        [
            ["khasras[].khasraId", "string", "opaque id", "Yes", "Khasra identifier"],
            ["khasras[].khasraNumber", "string", "parcel no", "Yes", "Khasra number as displayed"],
            ["khasras[].squareId", "string", "opaque id", "Yes", "Square identifier"],
            ["khasras[].squareNumber", "string", "display no", "Yes", "Square number"],
            ["khasras[].kilaId", "string", "opaque id", "Yes", "Kila identifier"],
            ["khasras[].kilaNumber", "string", "display no", "Yes", "Kila number"],
            ["khasras[].mauzaId", "string", "opaque id", "Yes", "Parent mauza"],
            ["page", "object", "—", "Yes", "Pagination metadata"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "data": {\n'
        '    "khasras": [\n'
        '      {\n'
        '        "khasraId": "K-9001",\n'
        '        "khasraNumber": "123/1",\n'
        '        "squareId": "SQ-12",\n'
        '        "squareNumber": "12",\n'
        '        "kilaId": "KL-03",\n'
        '        "kilaNumber": "3",\n'
        '        "mauzaId": "M-100"\n'
        '      }\n'
        '    ]\n'
        '  }\n'
        '}',
        [
            ["200", "DCVAL-0000", "Success", "Present Khasra list; on select call API-09 with khasraId, squareId, kilaId"],
            ["404", "ERR-4002", "Unknown mauzaId", "Refresh mauzas"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [["BR-05", "Khasras returned SHALL belong to the path mauzaId", "Service", "ERR-4002"]],
        [
            "Urban path selects a Khasra row then invokes API-09 (GetValuationMethod) with khasraId, squareId and kilaId.",
            "Field alias khasraNo MAY be accepted as synonym of khasraNumber for backward compatibility.",
        ],
    )

    # API-06 classifications
    add_operation(
        "API-06", "Get Property Classifications",
        "Returns Rural property classifications: Residential, Commercial, Agriculture, Industrial.",
        "CAP-04", "GET", "/dc-valuation/v1/property-classifications", "dcval.read.masters",
        ["Caller authenticated.", "PERA user has selected property type RURAL."],
        ["No state change. Classification list returned."],
        [["propertyType", "query", "string", "M", "RURAL", "RURAL", "Shall be RURAL"]],
        None,
        "GET /dc-valuation/v1/property-classifications?propertyType=RURAL HTTP/1.1\n"
        "Authorization: Bearer {token}\n"
        "X-Correlation-Id: {uuid}",
        [
            ["classifications[].code", "string", "enum", "Yes", "RESIDENTIAL | COMMERCIAL | AGRICULTURE | INDUSTRIAL"],
            ["classifications[].name", "string", "UTF-8", "Yes", "Display name"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "data": {\n'
        '    "classifications": [\n'
        '      { "code": "RESIDENTIAL", "name": "Residential" },\n'
        '      { "code": "COMMERCIAL", "name": "Commercial" },\n'
        '      { "code": "AGRICULTURE", "name": "Agriculture" },\n'
        '      { "code": "INDUSTRIAL", "name": "Industrial" }\n'
        '    ]\n'
        '  }\n'
        '}',
        [
            ["200", "DCVAL-0000", "Success", "Bind classification control"],
            ["422", "ERR-4001", "propertyType not RURAL", "Use Urban path APIs"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [["BR-06", "Classifications apply only when propertyType=RURAL", "Service", "ERR-4001"]],
        ["Selected classification code is passed to API-07."],
    )

    # API-07 property areas
    add_operation(
        "API-07", "Get Property Areas by Tehsil",
        "Returns Property Areas under a Tehsil for the Rural path, filtered by property classification.",
        "CAP-04", "GET", "/dc-valuation/v1/tehsils/{tehsilId}/property-areas", "dcval.read.masters",
        ["Valid tehsilId.", "propertyType=RURAL.", "Valid classificationCode from API-06."],
        ["No state change. Property Area list returned."],
        [
            ["tehsilId", "path", "string", "M", "Must exist", "T-010", "Tehsil identifier"],
            ["propertyType", "query", "string", "M", "RURAL", "RURAL", "Shall be RURAL"],
            ["classificationCode", "query", "string", "M", "enum", "RESIDENTIAL", "From API-06"],
            ["page", "query", "integer", "O", "≥1", "1", "Page"],
            ["size", "query", "integer", "O", "1–200", "50", "Size"],
        ],
        None,
        "GET /dc-valuation/v1/tehsils/T-010/property-areas?propertyType=RURAL&classificationCode=RESIDENTIAL HTTP/1.1\n"
        "Authorization: Bearer {token}\n"
        "X-Correlation-Id: {uuid}",
        [
            ["propertyAreas[].propertyAreaId", "string", "opaque id", "Yes", "Property area identifier"],
            ["propertyAreas[].propertyAreaCode", "string", "code", "Yes", "Property area code"],
            ["propertyAreas[].propertyAreaName", "string", "UTF-8", "Yes", "Property area name"],
            ["propertyAreas[].classificationCode", "string", "enum", "Yes", "Classification applied"],
            ["propertyAreas[].tehsilId", "string", "opaque id", "Yes", "Parent tehsil"],
            ["page", "object", "—", "Yes", "Pagination metadata"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "data": {\n'
        '    "propertyAreas": [\n'
        '      {\n'
        '        "propertyAreaId": "PA-55",\n'
        '        "propertyAreaCode": "PA-55",\n'
        '        "propertyAreaName": "Sample Area",\n'
        '        "classificationCode": "RESIDENTIAL",\n'
        '        "tehsilId": "T-010"\n'
        '      }\n'
        '    ]\n'
        '  }\n'
        '}',
        [
            ["200", "DCVAL-0000", "Success", "Bind Property Area dropdown"],
            ["400", "ERR-1001", "Missing/invalid query params", "Correct and resubmit"],
            ["404", "ERR-4002", "Unknown tehsilId", "Refresh tehsils"],
            ["422", "ERR-4001", "Invalid classification or not RURAL", "Correct selection"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [["BR-07", "Property areas SHALL match tehsilId + classificationCode + RURAL", "Service", "ERR-4001"]],
        ["Selected propertyAreaId is input to API-08, API-10 (full-area valuation) and API-11 (khasra-based valuation)."],
    )

    # API-08
    add_operation(
        "API-08", "Get Khasra / Kila Nos by Property Area",
        "Returns Khasra and/or Kila lists for the selected Rural Property Area. PERA MAY allow multi-select of khasras.",
        "CAP-04", "GET", "/dc-valuation/v1/property-areas/{propertyAreaId}/khasras", "dcval.read.masters",
        ["Valid propertyAreaId from API-07.", "Rural path in progress."],
        ["No state change. Khasra and Kila identifier lists returned for selection."],
        [
            ["propertyAreaId", "path", "string", "M", "Must exist", "PA-55", "Property area identifier"],
            ["listType", "query", "string", "O", "KHASRA|KILA|BOTH", "BOTH", "Which list(s) to return"],
            ["page", "query", "integer", "O", "≥1", "1", "Page"],
            ["size", "query", "integer", "O", "1–200", "50", "Size"],
        ],
        None,
        "GET /dc-valuation/v1/property-areas/PA-55/khasras?listType=BOTH&page=1&size=50 HTTP/1.1\n"
        "Authorization: Bearer {token}\n"
        "X-Correlation-Id: {uuid}",
        [
            ["khasras[].khasraId", "string", "opaque id", "Yes*", "Khasra identifier (*when list includes khasras)"],
            ["khasras[].khasraNumber", "string", "parcel no", "Yes*", "Khasra number"],
            ["khasras[].squareId", "string", "opaque id", "No", "Square identifier if applicable"],
            ["khasras[].squareNumber", "string", "display no", "No", "Square number if applicable"],
            ["khasras[].kilaId", "string", "opaque id", "No", "Kila identifier if applicable"],
            ["khasras[].kilaNumber", "string", "display no", "No", "Kila number if applicable"],
            ["kilas[].kilaId", "string", "opaque id", "Yes*", "Kila identifier (*when list includes kilas)"],
            ["kilas[].kilaNumber", "string", "display no", "Yes*", "Kila number"],
            ["propertyAreaId", "string", "opaque id", "Yes", "Parent property area"],
            ["page", "object", "—", "Yes", "Pagination metadata"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "data": {\n'
        '    "propertyAreaId": "PA-55",\n'
        '    "khasras": [\n'
        '      {\n'
        '        "khasraId": "K-7001",\n'
        '        "khasraNumber": "45",\n'
        '        "squareId": "SQ-01",\n'
        '        "squareNumber": "1",\n'
        '        "kilaId": "KL-01",\n'
        '        "kilaNumber": "1"\n'
        '      }\n'
        '    ],\n'
        '    "kilas": [\n'
        '      { "kilaId": "KL-01", "kilaNumber": "1" },\n'
        '      { "kilaId": "KL-02", "kilaNumber": "2" }\n'
        '    ]\n'
        '  }\n'
        '}',
        [
            ["200", "DCVAL-0000", "Success", "Allow single or multi Khasra selection; then call API-11"],
            ["404", "ERR-4002", "Unknown propertyAreaId", "Refresh property areas"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [["BR-08", "Khasras/Kilas returned SHALL belong to propertyAreaId", "Service", "ERR-4002"]],
        [
            "Where PERA selects Full Property Area valuation (no khasra selection), call API-10 instead of API-08/API-11.",
            "For multi-select Rural khasra valuation, PERA passes one or more khasra selections to API-11 (GetValuationMethod).",
        ],
    )

    # API-09 — Urban GetValuationMethod
    add_operation(
        "API-09", "Get Valuation by Khasra No (GetValuationMethod)",
        "Returns DC valuation details for a selected Urban Khasra context. "
        "Consumer passes khasraId, squareId and kilaId selected from API-05.",
        "CAP-05", "POST", "/dc-valuation/v1/valuations/by-khasra", "dcval.read.valuation",
        [
            "Valid khasraId, squareId and kilaId from API-05.",
            "Caller has dcval.read.valuation.",
        ],
        ["No state change. DC valuation details returned for the selected khasra context."],
        [],
        [
            ["khasraId", "string", "opaque id", "M", "Must exist", "K-9001"],
            ["squareId", "string", "opaque id", "M", "Must match khasra context", "SQ-12"],
            ["kilaId", "string", "opaque id", "M", "Must match khasra context", "KL-03"],
            ["asOfDate", "date", "YYYY-MM-DD", "O", "Optional effective date", "2026-08-10"],
        ],
        "POST /dc-valuation/v1/valuations/by-khasra HTTP/1.1\n"
        "Host: {uat-host}\n"
        "Authorization: Bearer {token}\n"
        "Content-Type: application/json; charset=utf-8\n"
        "X-Correlation-Id: {uuid}\n\n"
        "{\n"
        '  "khasraId": "K-9001",\n'
        '  "squareId": "SQ-12",\n'
        '  "kilaId": "KL-03"\n'
        "}",
        [
            ["khasraId", "string", "opaque id", "Yes", "Khasra identifier"],
            ["squareId", "string", "opaque id", "Yes", "Square identifier"],
            ["kilaId", "string", "opaque id", "Yes", "Kila identifier"],
            ["dcRate", "string", "decimal string", "Yes", "DC valuation rate amount"],
            ["structureRate", "string", "decimal string", "Yes", "Structure rate amount"],
            ["location", "string", "UTF-8", "Yes", "Location description / locality"],
            ["classification", "string", "UTF-8 / code", "Yes", "Property classification for the rate"],
            ["unitOfMeasurement", "string", "token", "Yes", "e.g. PER_MARLA / PER_KANAL / PER_SQFT"],
            ["currency", "string", "ISO code", "Yes", "PKR"],
            ["effectiveFrom", "date", "YYYY-MM-DD", "Yes", "Rate effective from"],
            ["scheduleReference", "string", "ref", "No", "PLRA schedule / notification reference"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "responseMessage": "DC valuation rate retrieved successfully",\n'
        '  "correlationId": "{uuid}",\n'
        '  "transactionId": "TXN-20260810-000090",\n'
        '  "timestamp": "2026-08-10T15:40:00+05:00",\n'
        '  "data": {\n'
        '    "khasraId": "K-9001",\n'
        '    "squareId": "SQ-12",\n'
        '    "kilaId": "KL-03",\n'
        '    "dcRate": "2500000.00",\n'
        '    "structureRate": "500000.00",\n'
        '    "location": "Sample Urban Locality",\n'
        '    "classification": "RESIDENTIAL",\n'
        '    "unitOfMeasurement": "PER_MARLA",\n'
        '    "currency": "PKR",\n'
        '    "effectiveFrom": "2026-01-01",\n'
        '    "scheduleReference": "DC-VAL-LHR-2026"\n'
        '  },\n'
        '  "errors": null\n'
        '}',
        [
            ["200", "DCVAL-0000", "Rate found", "Display rate fields in Activity Register"],
            ["404", "ERR-4003", "Khasra context not found or no rate published", "Inform user; do not invent rate"],
            ["422", "ERR-4005", "squareId/kilaId mismatch for khasraId", "Re-select khasra from API-05"],
            ["401", "ERR-2001", "Auth failure", "Re-authenticate"],
            ["403", "ERR-3001", "Missing valuation scope", "Escalate"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [
            ["BR-09", "If no published DC rate exists, return ERR-4003; never fabricate a rate", "Service", "ERR-4003"],
            ["BR-10", "khasraId + squareId + kilaId SHALL form a consistent Urban parcel key", "Service", "ERR-4005"],
            ["BR-10A", "Returned fields dcRate, structureRate, location, classification, unitOfMeasurement are mandatory when SUCCESS", "Service", "N/A"],
        ],
        [
            "Business alias: GetValuationMethod (single khasra / Urban path).",
            "Sample amounts are synthetic and non-authoritative.",
            "For Rural multi-Khasra selection, use API-11. For Rural full Property Area, use API-10.",
        ],
    )

    # API-10 — Rural full property area valuation
    add_operation(
        "API-10", "Get Valuation by Property Area",
        "Returns DC valuation details for a Rural Property Area when the citizen/user selects "
        "full property area valuation (no khasra multi-select).",
        "CAP-05", "GET", "/dc-valuation/v1/property-areas/{propertyAreaId}/valuation", "dcval.read.valuation",
        [
            "Valid propertyAreaId from API-07.",
            "Rural path; full-area valuation mode selected.",
            "Caller has dcval.read.valuation.",
        ],
        ["No state change. DC valuation details returned for the Property Area."],
        [
            ["propertyAreaId", "path", "string", "M", "Must exist", "PA-55", "Property area identifier"],
            ["asOfDate", "query", "date", "O", "YYYY-MM-DD", "2026-08-10", "Optional valuation effective date"],
        ],
        None,
        "GET /dc-valuation/v1/property-areas/PA-55/valuation HTTP/1.1\n"
        "Authorization: Bearer {token}\n"
        "X-Correlation-Id: {uuid}",
        [
            ["propertyAreaId", "string", "opaque id", "Yes", "Property area identifier"],
            ["propertyAreaName", "string", "UTF-8", "Yes", "Property area display name"],
            ["dcRate", "string", "decimal string", "Yes", "DC valuation rate amount"],
            ["structureRate", "string", "decimal string", "Yes", "Structure rate amount"],
            ["location", "string", "UTF-8", "Yes", "Location description"],
            ["classification", "string", "UTF-8 / code", "Yes", "Property classification"],
            ["unitOfMeasurement", "string", "token", "Yes", "Unit of measurement"],
            ["currency", "string", "ISO code", "Yes", "PKR"],
            ["effectiveFrom", "date", "YYYY-MM-DD", "Yes", "Rate effective from"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "responseMessage": "Property area DC valuation retrieved",\n'
        '  "correlationId": "{uuid}",\n'
        '  "transactionId": "TXN-20260810-000095",\n'
        '  "timestamp": "2026-08-10T15:42:00+05:00",\n'
        '  "data": {\n'
        '    "propertyAreaId": "PA-55",\n'
        '    "propertyAreaName": "Sample Rural Area",\n'
        '    "dcRate": "150000.00",\n'
        '    "structureRate": "25000.00",\n'
        '    "location": "Sample Rural Location",\n'
        '    "classification": "AGRICULTURE",\n'
        '    "unitOfMeasurement": "PER_KANAL",\n'
        '    "currency": "PKR",\n'
        '    "effectiveFrom": "2026-01-01"\n'
        '  },\n'
        '  "errors": null\n'
        '}',
        [
            ["200", "DCVAL-0000", "Rate found", "Display full-area rate in Activity Register"],
            ["404", "ERR-4003", "Property area not found or no rate published", "Inform user; do not invent rate"],
            ["403", "ERR-3001", "Missing scope", "Escalate"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry"],
        ],
        [
            ["BR-10B", "Full-area valuation SHALL NOT require khasra selection", "Service", "N/A"],
            ["BR-10C", "If no published area rate exists, return ERR-4003", "Service", "ERR-4003"],
        ],
        [
            "Use this operation when Rural user chooses Full Property Area.",
            "When Rural user instead selects khasra(s)/kila(s), use API-08 then API-11.",
        ],
    )

    # API-11 — Rural multi-khasra GetValuationMethod
    add_operation(
        "API-11", "Get Valuation by Property Area & Khasra No(s) (GetValuationMethod)",
        "Returns DC valuation details for one or more selected Rural Khasras under a Property Area. "
        "Supports multi-Khasra selection in a single call (GetValuationMethod).",
        "CAP-05", "POST", "/dc-valuation/v1/valuations/by-property-area", "dcval.read.valuation",
        [
            "Valid propertyAreaId from API-07.",
            "One or more khasra selections from API-08 belonging to that property area.",
            "Caller has dcval.read.valuation.",
        ],
        ["No state change. Array of DC valuation results returned (one entry per requested Khasra)."],
        [],
        [
            ["propertyAreaId", "string", "opaque id", "M", "Must exist", "PA-55"],
            ["khasras", "array[object]", "1–50 items", "M", "Each item identifies a khasra context", "see sample"],
            ["khasras[].khasraId", "string", "opaque id", "M", "Must belong to propertyAreaId", "K-7001"],
            ["khasras[].squareId", "string", "opaque id", "O/C", "Required where Square applies", "SQ-01"],
            ["khasras[].kilaId", "string", "opaque id", "O/C", "Required where Kila applies", "KL-01"],
            ["asOfDate", "date", "YYYY-MM-DD", "O", "Optional effective date", "2026-08-10"],
        ],
        "POST /dc-valuation/v1/valuations/by-property-area HTTP/1.1\n"
        "Host: {uat-host}\n"
        "Authorization: Bearer {token}\n"
        "Content-Type: application/json; charset=utf-8\n"
        "Accept: application/json\n"
        "X-Correlation-Id: {uuid}\n\n"
        "{\n"
        '  "propertyAreaId": "PA-55",\n'
        '  "khasras": [\n'
        '    { "khasraId": "K-7001", "squareId": "SQ-01", "kilaId": "KL-01" },\n'
        '    { "khasraId": "K-7002", "squareId": "SQ-01", "kilaId": "KL-02" }\n'
        "  ]\n"
        "}",
        [
            ["valuations[].khasraId", "string", "opaque id", "Yes", "Khasra identifier"],
            ["valuations[].squareId", "string", "opaque id", "No", "Square identifier"],
            ["valuations[].kilaId", "string", "opaque id", "No", "Kila identifier"],
            ["valuations[].propertyAreaId", "string", "opaque id", "Yes", "Property area"],
            ["valuations[].dcRate", "string", "decimal string", "Yes*", "DC rate (*null if NOT_FOUND)"],
            ["valuations[].structureRate", "string", "decimal string", "Yes*", "Structure rate"],
            ["valuations[].location", "string", "UTF-8", "Yes*", "Location"],
            ["valuations[].classification", "string", "UTF-8 / code", "Yes*", "Classification"],
            ["valuations[].unitOfMeasurement", "string", "token", "Yes*", "Unit of measurement"],
            ["valuations[].currency", "string", "PKR", "No", "Present when rate found"],
            ["valuations[].resultStatus", "string", "FOUND|NOT_FOUND", "Yes", "Per-khasra outcome"],
            ["valuations[].message", "string", "UTF-8", "No", "Per-khasra diagnostic if NOT_FOUND"],
        ],
        '{\n'
        '  "status": "SUCCESS",\n'
        '  "responseCode": "DCVAL-0000",\n'
        '  "responseMessage": "DC valuation rates retrieved",\n'
        '  "correlationId": "{uuid}",\n'
        '  "transactionId": "TXN-20260810-000100",\n'
        '  "timestamp": "2026-08-10T15:45:00+05:00",\n'
        '  "data": {\n'
        '    "propertyAreaId": "PA-55",\n'
        '    "valuations": [\n'
        '      {\n'
        '        "khasraId": "K-7001",\n'
        '        "squareId": "SQ-01",\n'
        '        "kilaId": "KL-01",\n'
        '        "propertyAreaId": "PA-55",\n'
        '        "dcRate": "180000.00",\n'
        '        "structureRate": "20000.00",\n'
        '        "location": "Sample Rural Location",\n'
        '        "classification": "AGRICULTURE",\n'
        '        "unitOfMeasurement": "PER_MARLA",\n'
        '        "currency": "PKR",\n'
        '        "resultStatus": "FOUND"\n'
        '      },\n'
        '      {\n'
        '        "khasraId": "K-7002",\n'
        '        "squareId": "SQ-01",\n'
        '        "kilaId": "KL-02",\n'
        '        "propertyAreaId": "PA-55",\n'
        '        "dcRate": null,\n'
        '        "structureRate": null,\n'
        '        "location": null,\n'
        '        "classification": null,\n'
        '        "unitOfMeasurement": null,\n'
        '        "resultStatus": "NOT_FOUND",\n'
        '        "message": "No published DC rate for khasra"\n'
        '      }\n'
        '    ]\n'
        '  },\n'
        '  "errors": null\n'
        '}',
        [
            ["200", "DCVAL-0000", "Request accepted; per-item statuses in data", "Map each khasra rate into Activity Register"],
            ["400", "ERR-1001", "Body/schema invalid", "Correct and resubmit"],
            ["404", "ERR-4002", "Unknown propertyAreaId", "Refresh property areas"],
            ["422", "ERR-4004", "khasraId not in property area", "Remove invalid ids and resubmit"],
            ["403", "ERR-3001", "Missing scope", "Escalate"],
            ["500", "ERR-5001", "Provider fault", "Bounded retry with same payload"],
        ],
        [
            ["BR-11", "Maximum 50 khasra selections per request", "Gateway/Service", "ERR-1001"],
            ["BR-12", "All khasraIds SHALL belong to propertyAreaId", "Service", "ERR-4004"],
            ["BR-13", "Partial success is allowed: overall SUCCESS with per-item NOT_FOUND", "Service", "N/A"],
            ["BR-14", "Duplicate khasra contexts in request are de-duplicated server-side", "Service", "N/A"],
            ["BR-15", "Multiple selected khasras SHALL be valued in one GetValuationMethod call", "Service", "N/A"],
        ],
        [
            "Business alias: GetValuationMethod (Rural multi-khasra path).",
            "POST is used only to carry the array body; the operation is read-only and idempotent.",
            "PERA SHALL display per-Khasra outcomes and SHALL NOT invent rates for NOT_FOUND items.",
        ],
    )

    page_break(doc)

    # ===================== SECTION 7 =====================
    heading(doc, "7. Data Model and Data Dictionary", 1)
    heading(doc, "7.1 Canonical Entities", 2)
    add_table(
        doc,
        ["Entity", "Business Definition", "Owning System", "Appears In"],
        [
            ["PropertyType", "Urban or Rural land context for valuation navigation", "PLRA DCVAL", "API-01"],
            ["District", "Administrative district of Punjab", "CLRMIS", "API-02, API-03"],
            ["Tehsil", "Administrative tehsil under a district", "CLRMIS", "API-03, API-04, API-07"],
            ["Mauza", "Revenue estate used on Urban path", "CLRMIS", "API-04, API-05, API-09"],
            ["PropertyClassification", "Rural land use class", "PLRA DCVAL", "API-06, API-07"],
            ["PropertyArea", "Rural area unit under tehsil + classification", "PLRA DCVAL", "API-07, API-08, API-10, API-11"],
            ["Khasra", "Land parcel number with optional Square/Kila context", "CLRMIS", "API-05, API-08, API-09, API-11"],
            ["Square", "Square subdivision associated with a khasra context", "CLRMIS", "API-05, API-08, API-09, API-11"],
            ["Kila", "Kila subdivision associated with a khasra context", "CLRMIS", "API-05, API-08, API-09, API-11"],
            ["DcValuation", "Published DC rate for khasra or property-area context", "PLRA DCVAL", "API-09, API-10, API-11"],
        ],
    )

    heading(doc, "7.2 Field-Level Data Dictionary", 2)
    add_table(
        doc,
        ["Field Name", "Entity", "Data Type", "Length", "Format / Pattern", "Nullable", "Classification", "Definition"],
        [
            ["districtId", "District", "string", "64", "opaque", "N", "Internal", "Stable district identifier"],
            ["districtName", "District", "string", "200", "UTF-8", "N", "Public", "District display name"],
            ["tehsilId", "Tehsil", "string", "64", "opaque", "N", "Internal", "Stable tehsil identifier"],
            ["mauzaId", "Mauza", "string", "64", "opaque", "N", "Internal", "Stable mauza identifier"],
            ["propertyAreaId", "PropertyArea", "string", "64", "opaque", "N", "Internal", "Stable property area identifier"],
            ["classificationCode", "PropertyClassification", "string", "32", "ENUM", "N", "Public", "RESIDENTIAL/COMMERCIAL/AGRICULTURE/INDUSTRIAL"],
            ["khasraId", "Khasra", "string", "64", "opaque", "N", "Internal", "Stable khasra identifier"],
            ["khasraNumber", "Khasra", "string", "50", "UTF-8", "N", "Internal", "Khasra number as published"],
            ["squareId", "Square", "string", "64", "opaque", "N", "Internal", "Square identifier"],
            ["squareNumber", "Square", "string", "50", "UTF-8", "N", "Internal", "Square number"],
            ["kilaId", "Kila", "string", "64", "opaque", "N", "Internal", "Kila identifier"],
            ["kilaNumber", "Kila", "string", "50", "UTF-8", "N", "Internal", "Kila number"],
            ["dcRate", "DcValuation", "string", "32", "decimal", "Y", "Sensitive", "DC valuation rate amount"],
            ["structureRate", "DcValuation", "string", "32", "decimal", "Y", "Sensitive", "Structure rate amount"],
            ["location", "DcValuation", "string", "250", "UTF-8", "Y", "Public", "Location / locality"],
            ["classification", "DcValuation", "string", "64", "UTF-8/code", "Y", "Public", "Classification for the published rate"],
            ["unitOfMeasurement", "DcValuation", "string", "32", "token", "Y", "Public", "Unit of measurement for the rate"],
            ["currency", "DcValuation", "string", "3", "ISO 4217", "Y", "Public", "Currency code (PKR)"],
            ["effectiveFrom", "DcValuation", "date", "10", "YYYY-MM-DD", "Y", "Public", "Rate effective start"],
            ["correlationId", "—", "string", "36", "UUID", "N", "Internal", "Consumer correlation key"],
        ],
    )

    heading(doc, "7.3 Reference Code Sets", 2)
    add_table(
        doc,
        ["Code Set", "Extensible", "Code", "Description", "Effective From"],
        [
            ["PROPERTY_TYPE", "No", "URBAN", "Urban land navigation path", "v1"],
            ["PROPERTY_TYPE", "No", "RURAL", "Rural land navigation path", "v1"],
            ["PROPERTY_CLASSIFICATION", "Yes", "RESIDENTIAL", "Residential classification", "v1"],
            ["PROPERTY_CLASSIFICATION", "Yes", "COMMERCIAL", "Commercial classification", "v1"],
            ["PROPERTY_CLASSIFICATION", "Yes", "AGRICULTURE", "Agriculture classification", "v1"],
            ["PROPERTY_CLASSIFICATION", "Yes", "INDUSTRIAL", "Industrial classification", "v1"],
            ["VALUATION_ITEM_STATUS", "No", "FOUND", "Rate returned for khasra", "v1"],
            ["VALUATION_ITEM_STATUS", "No", "NOT_FOUND", "No published rate for khasra", "v1"],
        ],
    )

    heading(doc, "7.4 Data Classification and Handling Obligations", 2)
    add_table(
        doc,
        ["Classification", "Definition", "Consumer Handling Obligation"],
        [
            ["Public", "Disclosure carries no adverse consequence", "No special handling"],
            ["Internal", "Intended for institutional use only", "Access control; no onward disclosure outside PERA Activity Register purpose"],
            ["Sensitive", "Disclosure would cause institutional or individual harm", "Encryption at rest, restricted access, access logging"],
            ["Personal", "Identifies a natural person", "Not intentionally exposed by these APIs; if encountered, apply DSA personal-data rules"],
        ],
    )
    heading(doc, "7.4.1 Purpose Limitation", 3)
    add_para(
        doc,
        "Data obtained through this interface SHALL be used exclusively for PERA Activity Register DC valuation "
        "lookup as stated in Section 1.4 and permitted by the DSA (R2). Onward disclosure, secondary analytics, "
        "or replication into systems not named in Section 2.1 requires prior written authorisation of PLRA data governance.",
    )
    heading(doc, "7.4.2 Masking and Minimisation", 3)
    add_table(
        doc,
        ["Field", "Response Context", "Masking Rule", "Rationale"],
        [
            ["dcRate", "Gateway access logs", "Redacted / not logged in full", "Limit sensitive financial/land valuation exposure in logs"],
            ["Authorization", "All logs", "Redacted", "Credential protection"],
        ],
    )

    page_break(doc)

    # ===================== SECTION 8 =====================
    heading(doc, "8. Transaction and State Management", 1)
    heading(doc, "8.1 Transaction Lifecycle", 2)
    add_para(
        doc,
        "All operations in this interface are read-only lookups. There is no binding commitment, reservation, "
        "or reversible business action at the Provider. The logical lifecycle of a lookup is:",
    )
    add_table(
        doc,
        ["State", "Meaning", "Entered By", "Permitted Next States", "Terminal"],
        [
            ["RECEIVED", "Request accepted at gateway", "Ingress", "AUTHENTICATED, REJECTED", "No"],
            ["AUTHENTICATED", "Token and scope validated", "Security enforcement", "COMPLETED, FAILED", "No"],
            ["COMPLETED", "Response returned successfully", "Service", "—", "Yes"],
            ["FAILED", "Documented error returned", "Service/Gateway", "—", "Yes"],
            ["REJECTED", "Auth/throttle/schema rejection", "Gateway", "—", "Yes"],
        ],
    )

    heading(doc, "8.2 Verification and Commitment Semantics", 2)
    add_para(
        doc,
        "Not Applicable for binding commitment — DC rate lookup is informational for Activity Register use. "
        "Returned rates reflect PLRA published schedules at query time and do not by themselves create a "
        "land transaction, mutation, or payment obligation inside PLRA.",
        size=10, italic=True,
    )

    heading(doc, "8.2.1 Two-Phase Commitment", 3)
    add_para(doc, "Not Applicable — the interface does not support reserve-then-commit.", size=10, italic=True)

    heading(doc, "8.3 Idempotency and Duplicate Suppression", 2)
    add_table(
        doc,
        ["Parameter", "Specification"],
        [
            ["Idempotency Mechanism", "Natural idempotency of read operations; X-Idempotency-Key not required for v1"],
            ["Behaviour on Repeat", "Same request yields equivalent current data (rates may change if schedule updates)"],
            ["Consumer Obligation", "Safe to retry reads after transient failure following Section 9.4"],
        ],
    )

    heading(doc, "8.4 Compensating Transactions", 2)
    add_para(doc, "Not Applicable — no state-changing operations; no reversal API.", size=10, italic=True)

    heading(doc, "8.5 Indeterminate Outcome Handling", 2)
    for i, t in enumerate([
        "On timeout or connection reset, PERA SHALL NOT assume a valuation was obtained.",
        "PERA MAY reissue the same read request (same parameters / body).",
        "PERA SHALL NOT write a DC rate into the Activity Register unless a SUCCESS response with FOUND/rate data was received.",
        "After repeated indeterminate outcomes, escalate per Section 11.4.",
    ], 1):
        add_para(doc, f"{i}. {t}", size=10, space_after=3)

    heading(doc, "8.6 Reconciliation", 2)
    add_table(
        doc,
        ["Parameter", "Specification"],
        [
            ["Reconciliation Mechanism", "Optional sample comparison of Activity Register stored rates vs re-query of API-09/API-10/API-11"],
            ["Frequency", "Monthly joint sample (recommended)"],
            ["Reconciliation Key", "khasraId + effectiveFrom / scheduleReference"],
            ["Authoritative Record", "PLRA DC Valuation source prevails"],
        ],
    )

    page_break(doc)

    # ===================== SECTION 9 =====================
    heading(doc, "9. Fault Management and Error Handling", 1)
    heading(doc, "9.1 Fault Classification", 2)
    add_table(
        doc,
        ["Class", "Origin", "Retryable", "Consumer Handling"],
        [
            ["Validation Fault", "Malformed request", "No", "Correct; do not retry unmodified"],
            ["Authentication Fault", "Invalid/expired credential", "Conditionally", "Re-authenticate once, then escalate"],
            ["Authorisation Fault", "Insufficient scope", "No", "Escalate"],
            ["Business Rule Fault", "Wrong path/type combination", "No", "Correct user flow"],
            ["Resource Fault", "Unknown id / no rate", "No", "Refresh lists; inform user"],
            ["Throttling Fault", "Quota exceeded", "Yes", "Honour Retry-After"],
            ["System / Dependency Fault", "Provider-side failure", "Yes", "Bounded exponential backoff"],
            ["Timeout", "No response", "Conditionally", "Retry read; do not store assumed rate"],
        ],
    )

    heading(doc, "9.2 HTTP Status Code Application", 2)
    add_table(
        doc,
        ["Status", "Applied When", "Response Body"],
        [
            ["200 OK", "Request processed and outcome returned", "Envelope with data"],
            ["400 Bad Request", "Syntactic or schema violation", "Envelope with errors"],
            ["401 Unauthorized", "Credential absent/malformed/expired", "Envelope with error"],
            ["403 Forbidden", "Authenticated but not entitled", "Envelope with error"],
            ["404 Not Found", "Referenced resource does not exist", "Envelope with error"],
            ["422 Unprocessable", "Valid syntax but business rule violated", "Envelope with error"],
            ["429 Too Many Requests", "Rate/quota exceeded", "Envelope; Retry-After set"],
            ["500 Internal Server Error", "Unhandled Provider fault", "Envelope; no internal detail"],
            ["503 Service Unavailable", "Maintenance or dependency outage", "Envelope; Retry-After set"],
        ],
    )

    heading(doc, "9.3 Error Code Register", 2)
    add_table(
        doc,
        ["Code", "HTTP", "Class", "Condition", "Consumer Action", "Retryable"],
        [
            ["DCVAL-0000", "200", "Success", "Successful processing", "Continue", "N/A"],
            ["ERR-1001", "400", "Validation", "Schema/parameter validation failure", "Correct and resubmit", "No"],
            ["ERR-1002", "400", "Validation", "Missing mandatory propertyType/classification", "Supply required filter", "No"],
            ["ERR-2001", "401", "Authentication", "Token missing/invalid", "Obtain new token", "Conditionally"],
            ["ERR-2002", "401", "Authentication", "Token expired", "Re-authenticate", "Conditionally"],
            ["ERR-3001", "403", "Authorisation", "Insufficient scope", "Escalate to PLRA", "No"],
            ["ERR-4001", "422", "Business Rule", "Urban/Rural path mismatch", "Follow correct path", "No"],
            ["ERR-4002", "404", "Resource", "Unknown district/tehsil/mauza/area id", "Refresh master data", "No"],
            ["ERR-4003", "404", "Resource", "No DC rate published for khasra", "Inform user", "No"],
            ["ERR-4004", "422", "Business Rule", "Khasra not in property area (API-11)", "Remove invalid ids", "No"],
            ["ERR-4005", "422", "Business Rule", "squareId/kilaId mismatch for khasraId (API-09)", "Re-select from API-05", "No"],
            ["ERR-5001", "500", "System", "Unhandled provider fault", "Bounded retry", "Yes"],
            ["ERR-6001", "429", "Throttling", "Rate/quota exceeded", "Backoff + Retry-After", "Yes"],
        ],
    )
    add_para(doc, "Table 9-1: Error code register. Complete register maintained at Appendix D.", size=9, italic=True, color=GRAY)

    heading(doc, "9.4 Retry and Backoff Policy", 2)
    add_table(
        doc,
        ["Parameter", "Specification"],
        [
            ["Retry Eligibility", "Only faults marked retryable in Section 9.3"],
            ["Maximum Attempts", "3 including initial attempt"],
            ["Backoff Algorithm", "Exponential with jitter"],
            ["Initial Interval", "500 ms"],
            ["Maximum Interval", "8000 ms"],
            ["Exhaustion Behaviour", "Suspend automated retries; alert PERA operations; do not store assumed rates"],
        ],
    )

    heading(doc, "9.5 Resilience Obligations of the Consumer", 2)
    for t in [
        "Apply connection and read timeouts no greater than values in Section 10.1.",
        "Implement a circuit breaker on sustained PLRA failure.",
        "Degrade gracefully in Activity Register (show service unavailable; do not block unrelated PERA functions).",
        "Never present an indeterminate timeout as a confirmed DC rate.",
    ]:
        bullet(doc, t)

    page_break(doc)

    # ===================== SECTION 10 =====================
    heading(doc, "10. Non-Functional Requirements", 1)
    heading(doc, "10.1 Performance", 2)
    add_table(
        doc,
        ["Metric", "Operation Class", "Target", "Measurement Point"],
        [
            ["Response time (median)", "Master data reads (API-01–08)", "≤ 800 ms", "Gateway ingress to egress"],
            ["Response time (95th percentile)", "Master data reads", "≤ 2000 ms", "Gateway ingress to egress"],
            ["Response time (95th percentile)", "Valuation reads (API-09–11)", "≤ 2000 ms", "Gateway ingress to egress"],
            ["Provider processing timeout", "All", "10 seconds", "Provider"],
            ["Recommended consumer timeout", "All", "15 seconds", "PERA client"],
        ],
    )

    heading(doc, "10.2 Capacity and Throughput", 2)
    add_table(
        doc,
        ["Parameter", "Committed Value", "Behaviour on Exceedance"],
        [
            ["Sustained throughput", "20 requests/second (initial; revisable)", "Throttled with 429"],
            ["Peak burst", "40 rps for ≤ 60 seconds", "Throttled with 429"],
            ["Daily quota", "As agreed in onboarding schedule", "Rejected with 429"],
            ["Concurrent connections", "50", "Queued or rejected"],
            ["Maximum request payload", "64 KB", "413"],
            ["Maximum khasra selections in API-11", "50", "400 ERR-1001"],
        ],
    )

    heading(doc, "10.3 Availability and Continuity", 2)
    add_table(
        doc,
        ["Parameter", "Specification"],
        [
            ["Service Availability Target", "99.0% monthly (initial); target uplift by joint agreement"],
            ["Measurement Basis", "Successful transaction ratio at gateway; agreed exclusions for planned maintenance"],
            ["Planned Maintenance Window", "As published by PLRA (typically off-peak); timezone PKT"],
            ["Maintenance Notification", "Minimum 72 hours via designated channel"],
            ["Emergency Maintenance", "Shortest practicable notice"],
            ["RTO / RPO", "Per PLRA enterprise DR policy — communicated in operations runbook"],
            ["Consumer Obligation During Failover", "Follow DNS/endpoint instructions issued by PLRA Ops; retry with backoff"],
        ],
    )

    heading(doc, "10.4 Scalability", 2)
    add_para(
        doc,
        "The service scales horizontally behind the PLRA API gateway. PERA SHALL forecast demand spikes "
        "(e.g. campaign-driven Activity Register load) at least 15 working days in advance so that PLRA can "
        "review quota and capacity. Capacity increases are executed under PLRA change management.",
    )

    heading(doc, "10.5 Logging, Monitoring and Observability", 2)
    add_table(
        doc,
        ["Aspect", "Provider (PLRA)", "Consumer (PERA)"],
        [
            ["Transaction logging", "correlationId, client id, operation, outcome, latency", "Persist correlationId/transactionId on Activity Register record"],
            ["Sensitive data in logs", "Masked per Section 7.4.2", "Masked; do not log bearer tokens"],
            ["Health endpoint", "Issued out of band", "Poll within agreed frequency only"],
            ["Alerting", "Availability, latency, error rate", "Client failure rate / timeout alerts"],
        ],
    )

    heading(doc, "10.6 Data Retention and Disposal", 2)
    add_table(
        doc,
        ["Data Category", "Retention Period", "Disposal Method", "Responsible Party"],
        [
            ["Gateway transaction / audit logs", "Per PLRA retention schedule (min 1 year)", "Secure disposal", "PLRA"],
            ["Activity Register stored rates", "Per PERA retention + DSA", "Per PERA policy", "PERA"],
            ["Cached master data at PERA", "≤ 24 hours unless otherwise agreed", "Invalidate/refresh", "PERA"],
        ],
    )

    heading(doc, "10.7 Interoperability and Portability", 2)
    for t in [
        "Interface definition published in OpenAPI form (Appendix A) sufficient for client generation.",
        "Open JSON/HTTPS standards used throughout.",
        "Reference code sets published for programmatic synchronisation.",
        "Contract transfers if PERA re-platforms; re-certification required.",
    ]:
        bullet(doc, t)

    page_break(doc)

    # ===================== SECTION 11 =====================
    heading(doc, "11. Service Levels and Operational Support", 1)
    heading(doc, "11.1 Service Level Commitments", 2)
    add_table(
        doc,
        ["SLA ID", "Commitment", "Target", "Measurement Window", "Reporting Frequency"],
        [
            ["SLA-01", "Service availability", "99.0%", "Calendar month", "Monthly"],
            ["SLA-02", "Response time at 95th percentile", "≤ 2000 ms", "Calendar month", "Monthly"],
            ["SLA-03", "Successful transaction ratio", "≥ 99.0% excluding client 4xx", "Calendar month", "Monthly"],
            ["SLA-04", "Planned maintenance notice", "≥ 72 hours", "Per event", "Per event"],
        ],
    )

    heading(doc, "11.2 Incident Severity Classification", 2)
    add_table(
        doc,
        ["Severity", "Definition", "Acknowledgement", "Update Cadence", "Target Resolution"],
        [
            ["S1 — Critical", "Complete unavailability of valuation APIs", "30 minutes", "Every 60 minutes", "8 hours"],
            ["S2 — High", "Major function unavailable (e.g. all valuation calls failing)", "1 hour", "Every 2 hours", "24 hours"],
            ["S3 — Medium", "Degraded performance or partial path failure with workaround", "4 hours", "Daily", "3 business days"],
            ["S4 — Low", "Minor defect, query or enhancement", "1 business day", "As agreed", "As prioritised"],
        ],
    )

    heading(doc, "11.3 Support Channels", 2)
    add_table(
        doc,
        ["Channel", "Detail", "Hours of Operation", "Applicable Severity"],
        [
            ["Service desk / ticketing", "PLRA integration ticketing (details on onboarding)", "Business hours + on-call for S1", "All"],
            ["Electronic mail", "Functional mailbox issued on onboarding", "Business hours", "S3, S4"],
            ["Telephone / emergency", "On-call arrangement issued on onboarding", "24×7 for S1", "S1, S2"],
        ],
    )
    add_para(
        doc,
        "All incidents SHALL be raised through the ticketing channel and quoted by correlation identifier. "
        "Verbal reports alone do not start the SLA clock.",
        size=10,
    )

    heading(doc, "11.4 Escalation Matrix", 2)
    add_table(
        doc,
        ["Level", "Elapsed Time", "Provider Contact", "Consumer Contact", "Authority"],
        [
            ["L1", "On raise", "PLRA API L1 Support", "PERA Integration L1", "Operational"],
            ["L2", "S1: 2h / S2: 4h", "PLRA API Operations Lead", "PERA Technical Lead", "Managerial"],
            ["L3", "S1: 4h / S2: 8h", "PLRA System Owner", "PERA IT Head", "Departmental"],
            ["L4", "S1: 8h", "Director IT PLRA", "PERA Executive Focal", "Steering"],
        ],
    )

    heading(doc, "11.5 Change Management", 2)
    add_table(
        doc,
        ["Change Type", "Definition", "Notice Period", "Consumer Retest Required", "Approval Authority"],
        [
            ["Standard", "Low-risk, no contract impact", "72 hours", "No", "PLRA Ops"],
            ["Normal", "Additive non-breaking contract change", "15 calendar days", "Regression recommended", "PLRA Architecture"],
            ["Major", "Breaking change / new major version", "180 calendar days", "Yes — full certification", "Director IT / Steering"],
            ["Emergency", "Security or availability driven", "Shortest practicable", "As assessed", "PLRA Security + Director IT"],
        ],
    )

    heading(doc, "11.6 Service Review", 2)
    add_para(
        doc,
        "A joint service review SHALL be convened monthly (or as agreed) to review SLA attainment, incident trends, "
        "capacity forecasts, open change requests and the risk register. Minutes and actions are tracked by PLRA API Operations.",
    )

    page_break(doc)

    # ===================== SECTION 12 =====================
    heading(doc, "12. Verification, Testing and Certification", 1)
    heading(doc, "12.1 Test Strategy", 2)
    add_table(
        doc,
        ["Test Level", "Executed By", "Environment", "Objective", "Evidence"],
        [
            ["Unit and component", "PERA", "PERA", "Client correctness", "Internal reports"],
            ["Interface connectivity", "Both", "Sandbox", "Reachability, TLS, authentication", "Log extract"],
            ["Functional integration", "PERA", "Sandbox", "Contract conformance per operation", "Executed test log"],
            ["Negative and boundary", "PERA", "Sandbox", "Fault handling per Section 9", "Executed test log"],
            ["Security", "PLRA / third party", "Staging", "Control effectiveness", "Assessment report"],
            ["Performance", "Both", "Staging", "Conformance to Section 10", "Performance report"],
            ["User acceptance", "PERA", "Sandbox/Staging", "Activity Register business outcome", "Signed UAT report"],
        ],
    )

    heading(doc, "12.2 Test Environment Provisioning", 2)
    add_table(
        doc,
        ["Item", "Specification"],
        [
            ["Access Request Procedure", "After DSA + registration; form and lead time issued by PLRA Ops"],
            ["Credentials Issued", "Sandbox client credentials; time-bound"],
            ["Test Data Provisioning", "Synthetic districts/tehsils/mauzas/khasras/rates — no production personal data"],
            ["Environment Availability", "Agreed UAT window"],
            ["Data Reset Cadence", "As published for sandbox"],
            ["Support During Testing", "PLRA integration support during business hours"],
        ],
    )

    heading(doc, "12.3 Mandatory Test Case Register", 2)
    add_table(
        doc,
        ["TC ID", "Category", "Scenario", "Expected Result", "Result", "Evidence Ref."],
        [
            ["TC-01", "Connectivity", "TLS from whitelisted address", "Handshake succeeds", "Pending", ""],
            ["TC-02", "Connectivity", "Call from non-whitelisted address", "Connection refused", "Pending", ""],
            ["TC-03", "Authentication", "Obtain token with valid credentials", "Token with expected claims", "Pending", ""],
            ["TC-04", "Authentication", "Invoke with expired token", "ERR-2002", "Pending", ""],
            ["TC-05", "Authorisation", "Call valuation API without valuation scope", "ERR-3001; no data", "Pending", ""],
            ["TC-06", "Functional", "Urban path end-to-end to API-09 (khasraId+squareId+kilaId)", "DCVAL-0000 with dcRate, structureRate, location, classification, unitOfMeasurement", "Pending", ""],
            ["TC-07", "Functional", "Rural full Property Area API-10", "DCVAL-0000 with valuation fields", "Pending", ""],
            ["TC-07A", "Functional", "Rural multi-khasra API-11 GetValuationMethod", "Per-item FOUND/NOT_FOUND with valuation fields", "Pending", ""],
            ["TC-08", "Validation", "Omit mandatory propertyType", "ERR-1002", "Pending", ""],
            ["TC-09", "Business Rule", "Call Mauza API with propertyType=RURAL", "ERR-4001", "Pending", ""],
            ["TC-10", "Resource", "Unknown districtId", "ERR-4002", "Pending", ""],
            ["TC-11", "Resource", "Khasra with no published rate", "ERR-4003 or NOT_FOUND", "Pending", ""],
            ["TC-12", "Business Rule", "API-11 khasra not in area", "ERR-4004", "Pending", ""],
            ["TC-12A", "Business Rule", "API-09 squareId/kilaId mismatch", "ERR-4005", "Pending", ""],
            ["TC-13", "Throttling", "Exceed rate ceiling", "429 + Retry-After honoured", "Pending", ""],
            ["TC-14", "Data", "Urdu names round-trip", "No corruption", "Pending", ""],
            ["TC-15", "Resilience", "Timeout then retry read", "No fabricated stored rate", "Pending", ""],
        ],
    )

    heading(doc, "12.4 Entry and Exit Criteria", 2)
    add_table(
        doc,
        ["Phase", "Entry Criteria", "Exit Criteria"],
        [
            ["Interface testing", "Connectivity + credentials + test data", "Connectivity/auth cases passed"],
            ["Functional testing", "Interface testing exited", "Mandatory functional/negative cases passed"],
            ["Certification", "Prior phases exited; defects closed/deferred", "Signed certification; production authorisation"],
        ],
    )

    heading(doc, "12.5 Defect Management", 2)
    add_table(
        doc,
        ["Severity", "Definition", "Owner", "Target Closure", "Blocks Certification"],
        [
            ["Critical", "Prevents core valuation lookup", "As assigned", "5 business days", "Yes"],
            ["Major", "Significant impairment; workaround exists", "As assigned", "10 business days", "Yes"],
            ["Minor", "Limited impact", "As assigned", "As prioritised", "No"],
            ["Cosmetic", "Presentational/documentation", "As assigned", "As prioritised", "No"],
        ],
    )

    heading(doc, "12.6 Certification Record", 2)
    add_table(
        doc,
        ["Certification Item", "Status", "Verified By", "Date"],
        [
            ["All mandatory test cases executed and passed", "Pending", "", ""],
            ["Outstanding defects closed or formally deferred", "Pending", "", ""],
            ["Security assessment completed with no open critical finding", "Pending", "", ""],
            ["Performance targets demonstrated", "Pending", "", ""],
            ["Operational readiness confirmed", "Pending", "", ""],
            ["Production authorisation granted", "Pending", "", ""],
        ],
    )

    page_break(doc)

    # ===================== SECTION 13 =====================
    heading(doc, "13. Consumer Implementation and Onboarding", 1)
    heading(doc, "13.1 Onboarding Sequence", 2)
    add_table(
        doc,
        ["Step", "Activity", "Responsible", "Prerequisite", "Output"],
        [
            ["1", "Execute DSA/MoU and obtain approvals", "Both", "—", "Signed instrument"],
            ["2", "Register PERA system and nominate focals", "PERA", "Step 1", "Registration record"],
            ["3", "Establish network connectivity/whitelisting", "Both", "Step 2", "Reachability evidence"],
            ["4", "Issue sandbox credentials/certificates", "PLRA", "Step 3", "Credential handover"],
            ["5", "Develop cascading UI + API client", "PERA", "Step 4", "Working client"],
            ["6", "Execute mandatory test register", "PERA", "Step 5", "Test evidence"],
            ["7", "Complete security assessment", "PLRA/Third party", "Step 6", "Assessment report"],
            ["8", "Obtain certification and production authorisation", "Both", "Step 7", "Certification record"],
            ["9", "Issue production credentials and enable access", "PLRA", "Step 8", "Production access"],
            ["10", "Controlled go-live and hypercare", "Both", "Step 9", "Go-live report"],
        ],
    )

    heading(doc, "13.2 Implementation Obligations of the Consumer", 2)
    for t in [
        "Generate/validate client against Appendix A OpenAPI definition.",
        "Externalise endpoints, credentials and timeouts into configuration.",
        "Store credentials in a secrets manager; never commit secrets to source control.",
        "Implement tolerant reading for additive response fields.",
        "Persist correlationId and transactionId on the Activity Register business record.",
        "Implement Urban and Rural cascading flows exactly as Section 2.2.1.",
        "For Rural full Property Area, call API-10. For Rural multi-select khasras, call API-11 (GetValuationMethod) and handle per-item FOUND/NOT_FOUND.",
        "Never invent or hard-code DC rates when APIs fail.",
    ]:
        bullet(doc, t)

    heading(doc, "13.3 Cutover Plan", 2)
    add_table(
        doc,
        ["Parameter", "Specification"],
        [
            ["Cutover Approach", "Phased enablement inside PERA Activity Register after certification"],
            ["Cutover Window", "Jointly scheduled"],
            ["Pre-Cutover Verification", "Production smoke: auth + one Urban + one Rural multi-khasra path"],
            ["Hypercare Period", "Minimum 10 business days enhanced support"],
            ["Success Criteria", "No S1/S2 open; ≥ 99% success on smoke set; business sign-off"],
        ],
    )

    heading(doc, "13.4 Rollback Plan", 2)
    add_table(
        doc,
        ["Parameter", "Specification"],
        [
            ["Rollback Trigger", "S1 unresolved beyond L3 threshold, or data integrity concern on displayed rates"],
            ["Decision Authority", "Joint PLRA Director IT / PERA executive focal"],
            ["Rollback Procedure", "Disable PERA feature flag / revoke production route; fall back to pre-integration process"],
            ["Data Reconciliation on Rollback", "Mark API-sourced rates in-flight; re-validate before reuse"],
            ["Communication Plan", "Notify focals within 30 minutes of rollback decision"],
        ],
    )

    page_break(doc)

    # ===================== SECTION 14 =====================
    heading(doc, "14. Lifecycle and Maintenance", 1)
    heading(doc, "14.1 Interface Lifecycle States", 2)
    add_table(
        doc,
        ["State", "Meaning", "Consumer Implication"],
        [
            ["Draft", "Under design; contract not stable", "Do not implement for production"],
            ["Released", "Baselined and supported", "Implement and operate normally"],
            ["Deprecated", "Superseded; supported for defined period", "Plan migration"],
            ["Sunset", "Withdrawal date announced", "Complete migration before date"],
            ["Withdrawn", "No longer served", "Invocation fails"],
        ],
    )

    heading(doc, "14.2 Enhancement and Change Request Procedure", 2)
    for i, t in enumerate([
        "Requesting party submits CR with business justification, affected operations and desired effective date.",
        "PLRA architecture performs impact assessment (contract, security, capacity, consumer impact).",
        "Change classified per Section 11.5 and routed for approval.",
        "On approval, schedule change, revise this document, baseline and reissue.",
        "Notify PERA per notice period; repeat certification for major changes.",
    ], 1):
        add_para(doc, f"{i}. {t}", size=10, space_after=3)

    heading(doc, "14.3 Document Maintenance", 2)
    add_table(
        doc,
        ["Aspect", "Specification"],
        [
            ["Document Owner", "Director IT / API Governance Custodian, PLRA"],
            ["Update Trigger", "Any change to contract, security control, SLA or operational procedure"],
            ["Scheduled Review", "Annually"],
            ["Baseline Procedure", "Approval, version increment, repository commit, controlled redistribution"],
            ["Consistency Obligation", "This document, OpenAPI artefact and deployed implementation verified mutually consistent at each baseline"],
        ],
    )

    heading(doc, "14.4 Knowledge Transfer and Continuity", 2)
    add_para(
        doc,
        "Each party SHALL nominate a primary and deputy focal person. Runbooks and this specification remain "
        "under Document Owner custody. On focal-person change, handover SHALL be completed within 10 business days "
        "and the distribution list in Section 0.5 updated.",
    )

    page_break(doc)

    # ===================== SECTION 15 =====================
    heading(doc, "15. Risk Register", 1)
    add_table(
        doc,
        ["Risk ID", "Description", "Category", "Likelihood", "Impact", "Mitigation", "Residual", "Owner"],
        [
            ["RSK-01", "Incomplete DSA delays onboarding", "Organisational", "3", "4", "Early legal engagement; sandbox gated on DSA", "Medium", "Both"],
            ["RSK-02", "Master data gaps (missing mauza/khasra/rates)", "Data", "3", "4", "Data readiness checklist before UAT; ERR-4003 handling in UI", "Medium", "PLRA"],
            ["RSK-03", "Consumer stores invented rates on timeout", "Operational", "2", "5", "Mandatory indeterminate-outcome rules; certification tests", "Low", "PERA"],
            ["RSK-04", "Credential leakage from client repos", "Security", "2", "5", "Secrets vault mandate; scanning; rotation", "Low", "PERA"],
            ["RSK-05", "Urban/Rural path confusion causes wrong API usage", "Technical", "3", "3", "Clear journey in §2.2.1; BR validation codes", "Low", "Both"],
            ["RSK-06", "Rate schedule updates not reflected in PERA cache", "Data", "3", "3", "Short cache TTL; no long-term offline rate sheets", "Medium", "PERA"],
        ],
    )
    heading(doc, "15.1 Assumption Invalidation Triggers", 2)
    add_para(
        doc,
        "Where any assumption in Section 2.4 is invalidated, the affected party SHALL notify the other within "
        "one working day and a joint impact assessment SHALL be convened before further development or operation "
        "proceeds on the affected path.",
    )

    page_break(doc)

    # ===================== APPENDICES =====================
    heading(doc, "Appendix A — Machine-Readable Interface Definition", 1)
    add_table(
        doc,
        ["Attribute", "Value"],
        [
            ["Specification Standard", "OpenAPI 3.0.x"],
            ["Filename", "plra-dcval-api-v1.yaml"],
            ["Version", "0.2"],
            ["Checksum", "To be generated on baseline of OpenAPI artefact"],
            ["Repository Location", "PLRA Controlled Integration Library"],
            ["Rendered Documentation", "Developer portal URL issued on onboarding (if published)"],
        ],
    )
    add_para(
        doc,
        "Where the OpenAPI artefact and this narrative diverge, the divergence is a defect and SHALL be raised "
        "against the Document Owner; neither artefact silently overrides the other.",
        size=10, italic=True,
    )

    heading(doc, "Appendix B — Sample Payload Library", 1)
    add_table(
        doc,
        ["Sample ID", "Operation", "Scenario", "File Reference"],
        [
            ["S-01", "API-02", "Nominal districts success", "samples/api02-success.json"],
            ["S-02", "API-03", "Unknown districtId", "samples/api03-notfound.json"],
            ["S-03", "API-09", "Urban valuation success (GetValuationMethod)", "samples/api09-success.json"],
            ["S-04", "API-10", "Rural full Property Area valuation", "samples/api10-area-success.json"],
            ["S-05", "API-11", "Rural multi-khasra partial NOT_FOUND (GetValuationMethod)", "samples/api11-partial.json"],
            ["S-05", "API-04", "propertyType mismatch business rule", "samples/api04-br-fail.json"],
        ],
    )
    add_para(doc, "All values are synthetic. No sample SHALL be derived from live data.", size=9, italic=True, color=GRAY)

    heading(doc, "Appendix C — Client Collection and Test Assets", 1)
    add_table(
        doc,
        ["Asset", "Format", "Version", "Location", "Custodian"],
        [
            ["API client collection", "Postman", "0.1", "Issued on sandbox onboarding", "PLRA API Owner"],
            ["Environment configuration", "JSON", "0.1", "Issued on sandbox onboarding", "PLRA API Owner"],
            ["Synthetic test dataset", "JSON/CSV", "0.1", "Sandbox", "PLRA"],
            ["Automated conformance suite", "TBD", "0.1", "TBD", "PLRA QA"],
        ],
    )
    add_para(doc, "Collections are distributed without credentials. PERA populates credential variables locally.", size=9, italic=True, color=GRAY)

    heading(doc, "Appendix D — Consolidated Error Code Register", 1)
    add_table(
        doc,
        ["Range", "Class", "Allocation Owner"],
        [
            ["DCVAL-0000", "Success", "PLRA API Owner"],
            ["1000–1999 (ERR-1xxx)", "Validation", "PLRA API Owner"],
            ["2000–2999 (ERR-2xxx)", "Authentication", "PLRA Security / IdP"],
            ["3000–3999 (ERR-3xxx)", "Authorisation", "PLRA Security"],
            ["4000–4999 (ERR-4xxx)", "Business rule / resource", "PLRA API Owner"],
            ["5000–5999 (ERR-5xxx)", "System and dependency", "PLRA Ops"],
            ["6000–6999 (ERR-6xxx)", "Throttling and quota", "PLRA Gateway Ops"],
        ],
    )

    heading(doc, "Appendix E — Requirements Traceability Matrix", 1)
    add_table(
        doc,
        ["Requirement ID", "Requirement Summary", "Capability", "Operation", "Test Case", "Verification Status"],
        [
            ["REQ-01", "Select Rural/Urban", "CAP-01", "API-01", "TC-06/TC-07", "Open"],
            ["REQ-02", "Cascade District → Tehsil", "CAP-02", "API-02, API-03", "TC-06", "Open"],
            ["REQ-03", "Urban Mauza → Khasra → Rate", "CAP-03, CAP-05", "API-04, API-05, API-09", "TC-06", "Open"],
            ["REQ-04", "Rural classification → area → khasras", "CAP-04", "API-06, API-07, API-08", "TC-07", "Open"],
            ["REQ-05", "Rural full Property Area DC rates", "CAP-05", "API-10", "TC-07", "Open"],
            ["REQ-05A", "Multi-khasra DC rates (GetValuationMethod)", "CAP-05", "API-11", "TC-07A", "Open"],
            ["REQ-06", "Secure G2G access", "—", "All", "TC-03–TC-05", "Open"],
        ],
    )

    heading(doc, "Appendix F — Review Checklist", 1)
    add_table(
        doc,
        ["#", "Check", "Reviewer", "Result", "Observation Ref."],
        [
            ["1", "Scope and boundary explicit", "Architecture", "Pending", ""],
            ["2", "Every operation traces to a capability", "Architecture", "Pending", ""],
            ["3", "Naming/method conventions consistent", "Technical", "Pending", ""],
            ["4", "Fields defined with type/obligation/validation", "Technical", "Pending", ""],
            ["5", "AuthN/AuthZ specified and testable", "Security", "Pending", ""],
            ["6", "No secrets in document", "Security", "Pending", ""],
            ["7", "Data classification and purpose limitation stated", "Data Governance", "Pending", ""],
            ["8", "Read semantics / indeterminate outcome specified", "Technical", "Pending", ""],
            ["9", "Error register complete with consumer actions", "Technical", "Pending", ""],
            ["10", "SLAs measurable", "Quality Assurance", "Pending", ""],
            ["11", "Mandatory tests cover Urban and Rural paths", "Quality Assurance", "Pending", ""],
            ["12", "OpenAPI artefact version-matched", "Technical", "Pending", ""],
            ["13", "Versioning/deprecation policy stated", "Architecture", "Pending", ""],
            ["14", "Risk register populated", "Quality Assurance", "Pending", ""],
            ["15", "Placeholders resolved; guidance removed", "Document Owner", "Pending", ""],
        ],
    )

    heading(doc, "Appendix G — Review Observation Log", 1)
    add_table(
        doc,
        ["Obs. ID", "Raised By", "Section", "Observation", "Disposition", "Closed On"],
        [
            ["OBS-01", "—", "—", "No observations recorded in draft 0.1", "—", "—"],
        ],
    )

    heading(doc, "Appendix H — Acceptance and Sign-Off", 1)
    add_para(
        doc,
        "By signing below, the parties confirm that this document constitutes the agreed technical contract "
        "governing the PLRA DC Valuation Rate interface for the PERA Activity Register, and that they accept "
        "the obligations assigned to them within it.",
    )
    add_table(
        doc,
        ["", "Service Provider", "Service Consumer"],
        [
            ["Organisation", "Punjab Land Records Authority (PLRA)", "PERA"],
            ["Name", "", ""],
            ["Designation", "", ""],
            ["Signature", "", ""],
            ["Date", "", ""],
            ["Official Seal", "", ""],
        ],
    )

    add_para(doc, "— End of Document —", size=12, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18)

    out = "/workspace/PLRA_PERA_DC_Valuation_API_Specification.docx"
    doc.save(out)
    print(f"Saved {out}")
    return out


if __name__ == "__main__":
    build()
