#!/usr/bin/env python3
"""Generate Sadia Ghafoor resume as a .docx file."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, size=10, bold=False, italic=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D0D0D0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(cell, text):
    p = cell.add_paragraph()
    set_paragraph_spacing(p, before=8, after=2, line=1.0)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True, color=RGBColor(0x33, 0x33, 0x33))
    add_horizontal_line(p)
    return p


def add_bullet(cell, text, bold_prefix=None, size=9.5):
    p = cell.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=1, line=1.1)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, size=size, bold=True)
        run = p.add_run(text)
        set_run_font(run, size=size)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size)
    return p


def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def clear_cell(cell):
    cell.text = ""


def build_resume():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # Name
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    run = p.add_run("Sadia Ghafoor")
    set_run_font(run, size=26, bold=True, color=RGBColor(0x11, 0x11, 0x11))

    # Title
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.0)
    run = p.add_run("Software Engineer")
    set_run_font(run, size=14, bold=True, color=RGBColor(0x66, 0x66, 0x66))

    # Summary
    summary_parts = [
        ("Detail-oriented and highly motivated SQA Engineer with ", False),
        ("1 year of hands-on experience", True),
        (" in software quality assurance and a strong academic foundation in Computer Science. Skilled in ", False),
        ("functional, regression, and UI testing", True),
        (" with practical exposure to test case design, defect tracking, and QA documentation. Proficient in tools such as ", False),
        ("Jira, Postman, and Selenium", True),
        (" (basic level), with a solid understanding of ", False),
        ("SDLC, STLC, and Agile methodologies", True),
        (". Adept at identifying bugs, improving product quality, and collaborating effectively with cross-functional teams. Recently graduated in Computer Science and eager to contribute to high-quality, user-focused software solutions.", False),
    ]
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=8, line=1.15)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in summary_parts:
        run = p.add_run(text)
        set_run_font(run, size=9.5, bold=bold, color=RGBColor(0x33, 0x33, 0x33))

    # Contact bar
    contact = doc.add_table(rows=1, cols=1)
    contact.autofit = True
    ccell = contact.cell(0, 0)
    clear_cell(ccell)
    shade_cell(ccell, "E8E8E8")
    set_cell_margins(ccell, top=50, bottom=50, left=80, right=80)
    cp = ccell.paragraphs[0]
    set_paragraph_spacing(cp, before=0, after=0, line=1.0)
    contacts = [
        ("sadiaghafoorsg007@gmail.com", True, "https://mailto:sadiaghafoorsg007@gmail.com"),
        ("  |  ", False, None),
        ("03143577288", False, None),
        ("  |  ", False, None),
        ("Lahore, Pakistan", False, None),
        ("  |  ", False, None),
        ("www.linkedin.com/in/sadia-ghafoor-sg007", True, "https://www.linkedin.com/in/sadia-ghafoor-sg007"),
    ]
    for text, is_link, url in contacts:
        run = cp.add_run(text)
        if is_link:
            set_run_font(run, size=9, color=RGBColor(0x1A, 0x73, 0xE8))
            run.underline = True
        else:
            set_run_font(run, size=9, color=RGBColor(0x33, 0x33, 0x33))

    # Make LinkedIn and email actual hyperlinks
    def add_hyperlink(paragraph, text, url, size=9):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1A73E8")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(size * 2)))
        rPr.append(szCs)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(rFonts)
        new_run.append(rPr)
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    # Rebuild contact paragraph with proper hyperlinks
    clear_cell(ccell)
    shade_cell(ccell, "E8E8E8")
    set_cell_margins(ccell, top=50, bottom=50, left=80, right=80)
    cp = ccell.paragraphs[0]
    set_paragraph_spacing(cp, before=0, after=0, line=1.0)
    add_hyperlink(cp, "sadiaghafoorsg007@gmail.com", "mailto:sadiaghafoorsg007@gmail.com")
    run = cp.add_run("  |  03143577288  |  Lahore, Pakistan  |  ")
    set_run_font(run, size=9, color=RGBColor(0x33, 0x33, 0x33))
    add_hyperlink(cp, "www.linkedin.com/in/sadia-ghafoor-sg007", "https://www.linkedin.com/in/sadia-ghafoor-sg007")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Two-column body
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.cell(0, 0), table.cell(0, 1)
    clear_cell(left)
    clear_cell(right)
    set_cell_margins(left, top=0, bottom=0, left=0, right=100)
    set_cell_margins(right, top=0, bottom=0, left=100, right=0)

    # Set column widths roughly 58% / 42%
    left.width = Cm(10.5)
    right.width = Cm(7.5)

    # LEFT: Education
    add_section_heading(left, "Education")
    p = left.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0, line=1.05)
    run = p.add_run("Bachelors in Computer Science")
    set_run_font(run, size=10.5, bold=True)

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=1.05)
    run = p.add_run("UMT - University of Management and Technology")
    set_run_font(run, size=10, bold=True)

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line=1.05)
    run = p.add_run("11/2021 - 07/2025")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x6B, 0x6B, 0x6B))
    run = p.add_run("\tLahore Pakistan")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x6B, 0x6B, 0x6B))

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=2, after=1, line=1.05)
    run = p.add_run("Courses")
    set_run_font(run, size=9.5, italic=True)

    for course in [
        "Object Oriented Programming (OOP)",
        "Database Management (DB)",
        "Data Structures",
        "Software Engineering",
        "Design and Analysis of Algorithm",
    ]:
        add_bullet(left, course, size=9)

    # LEFT: Work Experience
    add_section_heading(left, "Work Experience")

    # DinovaUX
    p = left.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0, line=1.05)
    run = p.add_run("Software Quality Assurance Engineer")
    set_run_font(run, size=10.5, bold=True)

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=1.05)
    run = p.add_run("DinovaUX")
    set_run_font(run, size=10, bold=True)

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line=1.05)
    run = p.add_run("02/2026 - 08/2026")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x6B, 0x6B, 0x6B))
    run = p.add_run("\tLahore")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x6B, 0x6B, 0x6B))

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line=1.1)
    run = p.add_run(
        "DinovaUX is a Lahore-based company delivering digital products and enterprise "
        "software solutions across learning, project management, monitoring, and "
        "organizational platforms."
    )
    set_run_font(run, size=9, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line=1.05)
    run = p.add_run("Achievements/Tasks:")
    set_run_font(run, size=9.5, bold=True, italic=True)

    dinova_tasks = [
        ("End-to-End QA: ", "Performed functional, regression, and UI testing across LMS, PHACT-PMO, Astrom, and Efficax-Al Munazam platforms"),
        ("Test Case Development: ", "Designed and executed test cases for course workflows, task management, alerts, approvals, and RBAC"),
        ("Defect Management: ", "Identified, documented, and tracked defects in Jira, collaborating with developers for timely resolution"),
        ("Workflow & SLA Validation: ", "Verified escalations, SLA rules, notifications, dashboards, and permission matrices"),
        ("API Testing: ", "Validated APIs and integrations using Postman to ensure reliable data flow between modules"),
        (None, "Participated in requirement reviews, sprint planning, and daily standups"),
    ]
    for prefix, text in dinova_tasks:
        add_bullet(left, text, bold_prefix=prefix)

    # Quick Stack
    p = left.add_paragraph()
    set_paragraph_spacing(p, before=8, after=0, line=1.05)
    run = p.add_run("Software Quality Assurance Engineer")
    set_run_font(run, size=10.5, bold=True)

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=1.05)
    run = p.add_run("Quick Stack Solutions")
    set_run_font(run, size=10, bold=True)

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line=1.05)
    run = p.add_run("06/2022 - 12/2022")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x6B, 0x6B, 0x6B))
    run = p.add_run("\tFaisalabad")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x6B, 0x6B, 0x6B))

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line=1.1)
    run = p.add_run(
        "QuickStack is a fast-growing startup focused on business growth consultancy and IT services."
    )
    set_run_font(run, size=9, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line=1.05)
    run = p.add_run("Achievements/Tasks:")
    set_run_font(run, size=9.5, bold=True, italic=True)

    qs_tasks = [
        ("Test Case Development: ", "Created and executed test cases for web applications, ensuring comprehensive coverage of functional and non-functional requirements"),
        ("Defect Management: ", "Identified, documented, and tracked defects, working closely with developers to resolve issues and improve product quality"),
        ("Bug Triage: ", "Prioritized and categorized defects based on severity and impact"),
        ("UI Testing: ", "Manually verified UI elements, ensuring alignment and adherence to design specifications"),
        ("Regression Bug Hunting: ", "Manually searched for regression bugs in previously tested areas after new changes were implemented"),
        (None, "Implemented and managed Software Development Life Cycle (SDLC) and Software Testing Life Cycle (STLC), streamlining workflows for efficiency"),
        (None, "Continuously identified gaps and improved testing methods and coverage"),
        (None, "Participated in requirement reviews, sprint planning, and daily standups"),
    ]
    for prefix, text in qs_tasks:
        add_bullet(left, text, bold_prefix=prefix)

    # RIGHT: Skills
    add_section_heading(right, "Skills")
    skills = [
        "Software Testing", "Test Planning", "SDLC & STLC", "Functional Testing",
        "Regression Testing", "UI Testing", "Defect Tracking", "Jira", "Postman",
        "Selenium (Basic)", "Agile / Scrum", "QA Documentation", "API Testing",
        "Test Case Design", "Load Testing", "Stress Testing", "Performance Testing",
    ]
    p = right.add_paragraph()
    set_paragraph_spacing(p, before=2, after=4, line=1.3)
    for i, skill in enumerate(skills):
        run = p.add_run(skill)
        set_run_font(run, size=9, bold=True, color=RGBColor(0x55, 0x55, 0x55))
        if i < len(skills) - 1:
            sep = p.add_run("  •  ")
            set_run_font(sep, size=9, color=RGBColor(0xAA, 0xAA, 0xAA))

    # RIGHT: Projects
    add_section_heading(right, "Projects")

    projects = [
        (
            "LMS (Learning Management System)",
            "(DinovaUX)",
            [
                "Online learning platform for course management, instructor training, learner enrollment, assessments, and certificates",
                "Validated self-paced/live course flows and role-based access for admins, instructors, and learners",
            ],
        ),
        (
            "PHACT-PMO",
            "(DinovaUX)",
            [
                "Project management system for tasks, entities, workflows, SLA rules, escalations, dashboards, and notifications",
                "Verified role-based access and end-to-end workflow behavior across user types",
            ],
        ),
        (
            "Astrom",
            "(DinovaUX)",
            [
                "Monitoring and automation platform for alerts, SOPs, and IPOM workflows",
                "Validated AI-assisted automation through Live Agent and alert/SOP reliability",
            ],
        ),
        (
            "Efficax-Al Munazam",
            "(DinovaUX)",
            [
                "Enterprise system for organizational obligations, task workflows, approvals, and notifications",
                "Verified user roles, approval chains, and task routing across processes",
            ],
        ),
        (
            "Foster Learner Management (Final Year Project)",
            "(10/2024 - 07/2025)",
            [
                "Implement personalized learning paths",
                "Integrate AI for automated grading and feedback",
                "Create a virtual classroom environment with real-time collaboration",
            ],
        ),
        (
            "Research Paper on Overleaf",
            "(05/2023 - 08/2023)",
            [
                "Research paper on Cloud computing Blockchain transportation system",
                "Explored how blockchain improves logistics through enhanced traceability, freight processes, and fraud reduction with transparent, immutable supply-chain records",
            ],
        ),
    ]

    for title, meta, bullets in projects:
        p = right.add_paragraph()
        set_paragraph_spacing(p, before=4, after=1, line=1.05)
        run = p.add_run(title + " ")
        set_run_font(run, size=10, bold=True)
        run = p.add_run(meta)
        set_run_font(run, size=9, italic=True, color=RGBColor(0x6B, 0x6B, 0x6B))
        for b in bullets:
            add_bullet(right, b, size=9)

    # RIGHT: Certificates
    add_section_heading(right, "Certificates")
    certs = [
        ("Technical Writing Competition", "(06/2019 - 07/2019)", "Winner in technical writing competition at Punjab University"),
        ("Calligraphy", "(03/2022 - 03/2022)", "Runner-up in calligraphy national level competition at Daira FAST"),
        ("Women Volleyball", "(03/2022 - 03/2022)", "Winner in Women Volleyball at Daira in FAST University FSD"),
        ("OPA Fiesta", "(06/2023 - 08/2024)", "Senior Team Member at Office of Participant Affairs in UMT"),
    ]
    for title, meta, desc in certs:
        p = right.add_paragraph()
        set_paragraph_spacing(p, before=3, after=0, line=1.05)
        run = p.add_run(title + " ")
        set_run_font(run, size=10, bold=True)
        run = p.add_run(meta)
        set_run_font(run, size=9, italic=True, color=RGBColor(0x6B, 0x6B, 0x6B))
        p = right.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.05)
        run = p.add_run(desc)
        set_run_font(run, size=9, color=RGBColor(0x55, 0x55, 0x55))

    # RIGHT: Languages
    add_section_heading(right, "Languages")
    p = right.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0, line=1.05)
    run = p.add_run("Urdu")
    set_run_font(run, size=10, bold=True)
    run = p.add_run("\t\tEnglish")
    set_run_font(run, size=10, bold=True)

    p = right.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=1.05)
    run = p.add_run("Native or Bilingual Proficiency")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x55, 0x55, 0x55))
    run = p.add_run("\tFull Professional Proficiency")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    out = "/workspace/Sadia_Ghafoor_Resume.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build_resume()
